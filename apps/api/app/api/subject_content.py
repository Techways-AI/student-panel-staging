from fastapi import APIRouter, Query, HTTPException, Response, Depends
from fastapi.responses import StreamingResponse, JSONResponse
import boto3
import json
import os
import io
import tempfile
import datetime
import logging
import re
import math
from dotenv import load_dotenv
from docx import Document
from app.api.auth import get_current_user
from app.models.user import User
from app.models.notes import GeneratedNotes
from app.models.subscription import Subscription
from app.models.quiz import Quiz
from app.models.content_library import ContentLibrary
from app.models.topic_mapping import TopicMapping
from sqlalchemy.orm import Session
from sqlalchemy import or_, func
from app.db.session import get_db
from app.utils.subscription_utils import should_lock_content, is_free_trial_expired, should_lock_content_by_semester, get_content_access_info
from app.services.analytics.posthog_client import capture_event, is_enabled

# Set up logger
logger = logging.getLogger(__name__)

# Import reportlab with fallback
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    print("WARNING: reportlab not available. PDF conversion will not work.")
    REPORTLAB_AVAILABLE = False

load_dotenv()  # Load .env file

# Create router without global authentication dependency
router = APIRouter()

# Get environment variables with validation
BUCKET = os.getenv('S3_BUCKET')
AWS_ACCESS_KEY_ID = os.getenv('AWS_ACCESS_KEY_ID')
AWS_SECRET_ACCESS_KEY = os.getenv('AWS_SECRET_ACCESS_KEY')
AWS_REGION = os.getenv('AWS_REGION')

# Debug: Print environment variables (mask the secret key)
print(f"DEBUG: S3_BUCKET = {BUCKET}")
print(f"DEBUG: AWS_ACCESS_KEY_ID = {'***' if AWS_ACCESS_KEY_ID else 'None'}")
print(f"DEBUG: AWS_SECRET_ACCESS_KEY = {'***' if AWS_SECRET_ACCESS_KEY else 'None'}")
print(f"DEBUG: AWS_REGION = {AWS_REGION}")

# Validate required environment variables
if not BUCKET:
    raise ValueError("S3_BUCKET environment variable is required")
if not AWS_ACCESS_KEY_ID:
    raise ValueError("AWS_ACCESS_KEY_ID environment variable is required")
if not AWS_SECRET_ACCESS_KEY:
    raise ValueError("AWS_SECRET_ACCESS_KEY environment variable is required")
if not AWS_REGION:
    raise ValueError("AWS_REGION environment variable is required")

import time
import threading


s3 = boto3.client(
    's3',
    aws_access_key_id=AWS_ACCESS_KEY_ID,
    aws_secret_access_key=AWS_SECRET_ACCESS_KEY,
    region_name=AWS_REGION,
    config=boto3.session.Config(
        connect_timeout=10,  # 10 seconds connection timeout
        read_timeout=30,     # 30 seconds read timeout
        retries={'max_attempts': 2}  # Retry up to 2 times
    )
)

def is_topic_completed_by_quiz(user_id: int, subject: str, unit: str, topic: str, db: Session) -> bool:
    """Check if a topic is completed based on quiz completion in the quiz table - with caching"""
    try:
        # Clean subject name to match database format
        from app.utils.subject_utils import clean_subject_name
        cleaned_subject = clean_subject_name(subject)
        
        # Checking completion (log removed)
        
        # Query database for completion status
        # Only consider quiz as completed if it has XP (meaning it was actually completed through proper quiz flow)
        # Quiz records created through daily goals or study plans don't have XP initially
        quiz = db.query(Quiz).filter(
            Quiz.student_id == user_id,
            Quiz.subject == cleaned_subject,
            Quiz.unit == unit,
            Quiz.topic == topic,
            Quiz.xp_topic.isnot(None)  # Only consider quizzes that have XP (actually completed)
        ).first()
        
        is_completed = quiz is not None and quiz.xp_topic is not None
        return is_completed
    except Exception as e:
        print(f"Error checking topic completion: {str(e)}")
        return False

def check_s3_file_exists(key: str):
    """Check if a file exists in S3 and return metadata"""
    try:
        response = s3.head_object(Bucket=BUCKET, Key=key)
        return {
            "exists": True,
            "size": response.get('ContentLength', 0),
            "last_modified": response.get('LastModified'),
            "content_type": response.get('ContentType')
        }
    except Exception as e:
        return {
            "exists": False,
            "error": str(e)
        }

def list_s3_files(prefix: str, max_keys: int = 100):
    """List files in S3 with a given prefix"""
    try:
        response = s3.list_objects_v2(
            Bucket=BUCKET,
            Prefix=prefix,
            MaxKeys=max_keys
        )
        files = []
        for obj in response.get('Contents', []):
            files.append({
                "key": obj['Key'],
                "size": obj['Size'],
                "last_modified": obj['LastModified']
            })
        return {
            "success": True,
            "files": files,
            "count": len(files)
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "files": []
        }

def normalize_path(path: str) -> str:
    """
    Normalize a path for consistent comparison.
    Removes extra spaces, normalizes separators, and handles common variations.
    """
    if not path:
        return ""
    
    # Strip whitespace and normalize
    normalized = path.strip()
    
    # Replace common unicode dash variants with a standard hyphen
    normalized = normalized.replace('–', '-').replace('—', '-').replace('−', '-')

    # Treat digit-digit separators uniformly: convert digit-hyphen-digit to digit_ digit
    # (e.g., 2-1 -> 2_1) so year/semester segments match metadata
    normalized = re.sub(r'(?<=\d)-(?=\d)', '_', normalized)

    # Replace multiple spaces with single space
    normalized = ' '.join(normalized.split())
    
    # Normalize path separators (ensure forward slashes)
    normalized = normalized.replace('\\', '/')
    
    # Remove trailing slashes
    normalized = normalized.rstrip('/')

    # Lowercase for case-insensitive matching (metadata paths are treated as case-insensitive)
    normalized = normalized.lower()

    # Split into parts for further normalization
    parts = [p for p in normalized.split('/') if p != '']

    # Map legacy prefix "bpharmacy" to "bpharma/pci" to match metadata convention
    if parts:
        if parts[0] == 'bpharmacy':
            parts[0] = 'bpharma'
            # Insert 'pci' if not already present as second segment
            if len(parts) < 2 or parts[1] != 'pci':
                parts.insert(1, 'pci')

    # Remove subject codes like "bp301t:" from segment starts
    cleaned_parts = []
    for seg in parts:
        cleaned_seg = re.sub(r'^[a-z0-9]+:\s*', '', seg)
        cleaned_parts.append(cleaned_seg)

    normalized = '/'.join(cleaned_parts)

    return normalized


def slugify(text: str) -> str:
    """
    Lightweight slugifier (kept inline to avoid new deps):
    'Benzene and Its Derivatives' -> 'benzene-and-its-derivatives'
    """
    text = text.lower().strip()
    text = re.sub(r"[^\w\s-]", "", text)
    return re.sub(r"[\s]+", "-", text)

@router.get('/get-document-id')
def get_document_id(
    fullPath: str = Query(None, description="Legacy full path lookup (kept for backward compatibility)"),
    topic: str = Query(None, description="Topic name to resolve via mapping → PCI → document"),
    university: str = Query("PCI", description="University name used for topic mapping"),
    db: Session = Depends(get_db)
):
    """
    Resolve a document id for AI/notes.

    Order of resolution (non-breaking):
    1) Topic + university mapping → PCI slug → s3_key → document_id
    2) Fallback: legacy folder-path lookup using fullPath (unchanged)
    """
    try:
        # ---------- 1) Topic-based resolution (PCI is canonical) ----------
        if topic:
            uni = (university or "PCI").strip()
            normalized_topic = topic.strip()
            pci_slug = None
            pci_topic = None
            mapping_used = False

            # Non-PCI → map to PCI
            if uni.lower() != "pci":
                mapping = (
                    db.query(TopicMapping)
                    .filter(
                        TopicMapping.university_name == uni,
                        func.lower(TopicMapping.university_topic) == normalized_topic.lower()
                    )
                    .first()
                )
                if mapping:
                    mapping_used = True
                    pci_slug = mapping.topic_slug
                    pci_topic = mapping.pci_topic
                else:
                    # No mapping; we'll fall back to legacy path if provided
                    logger.info(f"No topic mapping found for '{normalized_topic}' in university '{uni}'")
            else:
                # PCI direct
                pci_slug = slugify(normalized_topic)
                pci_topic = normalized_topic

            # If we resolved a slug, try to find an S3 key from content library (documents/notes preferred)
            if pci_slug:
                preferred_types = ["document", "notes"]
                def pick_entry(slug_value: str):
                    return (
                        db.query(ContentLibrary)
                        .filter(
                            ContentLibrary.topic_slug == slug_value,
                            ContentLibrary.file_type.in_(preferred_types)
                        )
                        .order_by(ContentLibrary.created_at.desc())
                        .first()
                    )

                def pick_entry_relaxed(slug_value: str):
                    return (
                        db.query(ContentLibrary)
                        .filter(
                            ContentLibrary.topic_slug.ilike(f"%{slug_value}%"),
                            ContentLibrary.file_type.in_(preferred_types)
                        )
                        .order_by(ContentLibrary.created_at.desc())
                        .first()
                    )

                doc_entry = pick_entry(pci_slug)

                # Fallback 1: broader match on slug fragment (handles mismatched mapping slugs)
                if not doc_entry:
                    doc_entry = pick_entry_relaxed(pci_slug)

                # Fallback 2: try slugified topic name if provided
                if not doc_entry and pci_topic:
                    derived_slug = slugify(pci_topic)
                    doc_entry = pick_entry(derived_slug) or pick_entry_relaxed(derived_slug)

                # Fallback 3: any file_type if still not found
                if not doc_entry:
                    doc_entry = (
                        db.query(ContentLibrary)
                        .filter(ContentLibrary.topic_slug == pci_slug)
                        .order_by(ContentLibrary.created_at.desc())
                        .first()
                    )

                if doc_entry:
                    s3_key = doc_entry.s3_key
                    # Reuse existing key-based resolver to keep behavior consistent
                    key_lookup = get_document_id_from_key(s3_key=s3_key)
                    if key_lookup and key_lookup.get("document_id"):
                        return {
                            "success": True,
                            "document_id": key_lookup.get("document_id"),
                            "file_key": s3_key,
                            "topic": normalized_topic,
                            "pci_topic": pci_topic,
                            "pci_slug": pci_slug,
                            "university": uni,
                            "source": "topic_mapping" if mapping_used else "pci_direct",
                        }

            # If topic path failed and no legacy fullPath provided, return graceful miss
            if not fullPath:
                return {
                    "success": False,
                    "document_id": None,
                    "topic": topic,
                    "university": uni,
                    "message": "No document found via topic mapping or PCI lookup",
                }

        # ---------- 2) Legacy folder-path lookup (unchanged behavior) ----------
        if not fullPath:
            return {
                "document_id": None,
                "success": False,
                "message": "fullPath is required when topic-based resolution fails or is not provided"
            }

        # Use the correct key for the metadata file (plural: documents.json)
        metadata_key = "metadata/documents.json"

        # Get the metadata file from S3 with timeout
        try:
            metadata_obj = s3.get_object(Bucket=BUCKET, Key=metadata_key)
            metadata_data = json.loads(metadata_obj['Body'].read())
        except Exception as s3_error:
            return {
                "document_id": None,
                "fullPath": fullPath,
                "error": f"Failed to load metadata: {str(s3_error)}"
            }

        # Strip filename from the fullPath to match the folderStructure.fullPath format
        path_parts = fullPath.split('/')
        if len(path_parts) > 0 and '.' in path_parts[-1]:  # If last part has a file extension
            folder_path = '/'.join(path_parts[:-1])
        else:
            folder_path = fullPath

        # Normalize the search path
        normalized_search_path = normalize_path(folder_path)

        document_id = None
        potential_matches = []

        for document in metadata_data:
            doc_path = document.get('folderStructure', {}).get('fullPath')
            if doc_path:
                normalized_doc_path = normalize_path(doc_path)

                if normalized_doc_path == normalized_search_path:
                    document_id = document.get('id')
                    break

                if normalized_search_path in normalized_doc_path or normalized_doc_path in normalized_search_path:
                    potential_matches.append({
                        'id': document.get('id'),
                        'path': doc_path,
                        'normalized': normalized_doc_path
                    })

        if document_id:
            return {
                "document_id": document_id,
                "fullPath": fullPath,
                "folder_path": folder_path,
                "normalized_path": normalized_search_path,
                "success": True
            }
        else:
            return {
                "document_id": None,
                "fullPath": fullPath,
                "folder_path": folder_path,
                "normalized_path": normalized_search_path,
                "success": False,
                "message": "No document found for the specified path",
                "potential_matches": potential_matches[:5] if potential_matches else [],
                "suggestions": [
                    "Check if the document has been uploaded to S3",
                    "Verify the path format matches the metadata structure",
                    "Ensure the document is properly indexed in the metadata",
                    "Check for path variations (spaces, case sensitivity, etc.)"
                ]
            }

    except Exception as e:
        return {
            "document_id": None,
            "fullPath": fullPath,
            "error": f"Failed to fetch document ID: {str(e)}",
            "success": False
        }


@router.get('/get-document-id-from-key')
def get_document_id_from_key(
    s3_key: str = Query(..., description="Exact S3 key / file_key to look up"),
):
    """
    Fetch document ID from metadata based on an S3 key (file_key or s3_key).

    This keeps existing behavior intact and simply adds a reverse lookup
    alongside the folder-path based lookup above.
    """
    try:
        if not s3_key:
            return {
                "document_id": None,
                "file_key": s3_key,
                "success": False,
                "message": "s3_key is required"
            }

        metadata_key = "metadata/documents.json"

        try:
            metadata_obj = s3.get_object(Bucket=BUCKET, Key=metadata_key)
            metadata_data = json.loads(metadata_obj['Body'].read())
        except Exception as s3_error:
            return {
                "document_id": None,
                "file_key": s3_key,
                "error": f"Failed to load metadata: {str(s3_error)}",
                "success": False
            }

        match = None
        potential_matches = []

        for document in metadata_data:
            doc_file_key = document.get('file_key') or document.get('s3_key')
            if not doc_file_key:
                continue

            if doc_file_key == s3_key:
                match = {
                    "document_id": document.get('id'),
                    "file_key": doc_file_key,
                    "name": document.get('name'),
                    "subject": document.get('subject'),
                    "unit": document.get('unit'),
                    "topic": document.get('topic'),
                    "folderStructure": document.get('folderStructure', {})
                }
                break

            if s3_key in doc_file_key or doc_file_key in s3_key:
                potential_matches.append({
                    "document_id": document.get('id'),
                    "file_key": doc_file_key,
                    "path": document.get('folderStructure', {}).get('fullPath')
                })

        if match:
            return {
                **match,
                "success": True
            }

        return {
            "document_id": None,
            "file_key": s3_key,
            "success": False,
            "message": "No document found for the specified S3 key",
            "potential_matches": potential_matches[:5] if potential_matches else [],
            "suggestions": [
                "Verify the file_key/s3_key exists in metadata/documents.json",
                "Ensure the key matches exactly (case-sensitive)",
                "Check if the document has been uploaded and indexed"
            ]
        }

    except Exception as e:
        return {
            "document_id": None,
            "file_key": s3_key,
            "error": f"Failed to fetch document ID: {str(e)}",
            "success": False
        }

@router.get('/get-file-key-from-document-id')
def get_file_key_from_document_id(document_id: str = Query(..., description="Document ID to search for")):
    """Fetch file key from metadata file based on document ID"""
    try:
        
        # Use the correct key for the metadata file (plural: documents.json)
        metadata_key = "metadata/documents.json"
        
        # Get the metadata file from S3 with timeout
        try:
            metadata_obj = s3.get_object(Bucket=BUCKET, Key=metadata_key)
            metadata_data = json.loads(metadata_obj['Body'].read())
        except Exception as s3_error:
            return {
                "file_key": None,
                "document_id": document_id,
                "error": f"Failed to load metadata: {str(s3_error)}",
                "success": False
            }
        
        file_key = None
        document_info = None
        match_count = 0
        
        for document in metadata_data:
            doc_id = document.get('id')
            if doc_id == document_id:
                # Get the file key from the document metadata
                file_key = document.get('file_key') or document.get('s3_key')
                document_info = {
                    'id': document.get('id'),
                    'name': document.get('name'),
                    'subject': document.get('subject'),
                    'unit': document.get('unit'),
                    'topic': document.get('topic'),
                    'folderStructure': document.get('folderStructure', {})
                }
                match_count += 1
                break
        
        if file_key:
            return {
                "file_key": file_key,
                "document_info": document_info,
                "document_id": document_id,
                "success": True
            }
        else:
            return {
                "file_key": None,
                "document_id": document_id,
                "success": False,
                "message": "No file key found for the specified document ID",
                "suggestions": [
                    "Check if the document has been properly uploaded to S3",
                    "Verify the document ID exists in the metadata",
                    "Ensure the document has a valid file_key or s3_key field"
                ]
            }
            
    except Exception as e:
        print(f"ERROR: Exception in get_file_key_from_document_id for document_id {document_id}: {str(e)}")
        return {
            "file_key": None,
            "document_id": document_id,
            "error": f"Failed to fetch file key: {str(e)}",
            "success": False
        }

def convert_docx_to_pdf(docx_content):
    """Convert DOCX content to PDF"""
    if not REPORTLAB_AVAILABLE:
        raise Exception("PDF conversion is not available. reportlab library is not installed.")
    
    try:
        # Create a temporary file to store the DOCX content
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            temp_docx.write(docx_content)
            temp_docx_path = temp_docx.name
        
        # Load the DOCX document
        doc = Document(temp_docx_path)
        
        # Create PDF buffer
        pdf_buffer = io.BytesIO()
        pdf_doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
        
        # Get styles
        styles = getSampleStyleSheet()
        normal_style = styles['Normal']
        heading_style = styles['Heading1']
        
        # Build PDF content
        story = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Determine if it's a heading (simple heuristic)
                if paragraph.style.name.startswith('Heading') or len(paragraph.text) < 100:
                    p = Paragraph(paragraph.text, heading_style)
                else:
                    p = Paragraph(paragraph.text, normal_style)
                story.append(p)
                story.append(Spacer(1, 12))
        
        # Build PDF
        pdf_doc.build(story)
        
        # Clean up temporary file
        os.unlink(temp_docx_path)
        
        # Get PDF content
        pdf_content = pdf_buffer.getvalue()
        pdf_buffer.close()
        
        return pdf_content
        
    except Exception as e:
        print(f"Error converting DOCX to PDF: {str(e)}")
        raise e





@router.get('/')
def get_subject_content(
    courseName: str, 
    yearSemester: str, 
    subjectName: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    try:
        # Main prefix for documents: bpharmacy/1-1/PS101: Human Anatomy and Physiology I/
        prefix = f'{courseName}/{yearSemester}/{subjectName}/'
        
        # Subject content requested (logs removed to reduce noise)
        
        # Get user's active subscription
        subscription = db.query(Subscription).filter(
            Subscription.user_id == current_user.id,
            Subscription.status == "active",
            Subscription.end_date > datetime.datetime.utcnow()
        ).first()
        
        # Get comprehensive access information
        access_info = get_content_access_info(current_user, yearSemester, subscription)
        
        # DISABLED: No longer using S3 - return empty structure
        # Use /api/curriculum/subject-content endpoint instead for subject structure
        content_structure = []
        
        result = []
        
        # Check subject-specific access for subject-based subscriptions
        subject_code = None
        has_subject_access_result = False
        if ':' in subjectName:
            subject_code = subjectName.split(':')[0].strip()
            
            # Check if user has access to this specific subject
            from app.utils.subscription_utils import has_subject_access
            has_subject_access_result = has_subject_access(current_user, subject_code, db)
        
        # Determine content access based on semester and subscription
        content_locked = should_lock_content_by_semester(current_user, yearSemester, subscription)
        
        # For subject-based subscriptions, apply proper access control
        if current_user.subscription_status == 'subject_based':
            if subject_code and not has_subject_access_result:
                # User doesn't have access to this subject - check if it's in current semester for 25% access
                user_year = current_user.year or 1
                user_semester = current_user.semester or 1
                
                # Get subject details from hardcoded data
                from app.api.subject_subscriptions import SUBJECTS_DATA
                subject_info = next(
                    ((code, name, year, semester) for code, name, year, semester in SUBJECTS_DATA 
                    if code == subject_code), None)
                
                if subject_info:
                    _, _, subject_year, subject_semester = subject_info
                    if subject_year == user_year and subject_semester == user_semester:
                        # Allow 25% free trial access to subjects in current semester
                        content_locked = False
                    else:
                        # Block access to subjects from different semesters
                        content_locked = True
                else:
                    # Subject not found in data - block access
                    content_locked = True
            else:
                # User has access to this subject - full access
                content_locked = False
        
        # User is considered "free" for free trial access based on subject access
        if current_user.subscription_status == 'subject_based':
            # Subject-based users get 25% access to non-purchased subjects in current semester
            is_free_user = not has_subject_access_result
        else:
            has_regular_subscription = subscription is not None and subscription.is_active
            is_free_user = not has_regular_subscription  # Regular users without paid plan use 25% free access
        trial_expired = is_free_trial_expired(current_user)
        
        # Access control determined (logs removed)
        
        # ULTRA-OPTIMIZED: Batch quiz completion check with enhanced XP caching
        from app.utils.subject_utils import clean_subject_name
        cleaned_subject = clean_subject_name(subjectName)
        
        # Get ALL quiz completions for this subject in one ultra-optimized query
        batch_start = time.time()
        
        # ULTRA-OPTIMIZED: Use specific columns and efficient filtering with index hints
        completed_quizzes = db.query(
            Quiz.unit, 
            Quiz.topic
        ).filter(
            Quiz.student_id == current_user.id,
            Quiz.subject == cleaned_subject,
            Quiz.xp_topic.isnot(None)  # Only completed quizzes with XP
        ).distinct().all()
        
        # Create lookup dict for O(1) completion checks
        completion_lookup = {}
        for quiz in completed_quizzes:
            key = f"{quiz.unit}_{quiz.topic}"
            completion_lookup[key] = True
            
        batch_time = time.time() - batch_start
        
        # Calculate free access limits (25% of topics) for free users with active trial
        total_topics_count = sum(len(unit_data.get('topics', [])) for unit_data in content_structure)
        free_access_limit = 0
        free_topics_granted = 0
        if is_free_user and not content_locked and total_topics_count > 0:
            free_access_limit = max(1, math.ceil(total_topics_count * 0.25))

        # OPTIMIZED: Process the pre-loaded content structure instead of making S3 calls
        for unit_index, unit_data in enumerate(content_structure):
            unit_name = unit_data['unitName']
            topic_list = []
            
            for topic_index, topic_data in enumerate(unit_data['topics']):
                topic_name = topic_data['topicName']
                files = topic_data['files']
                video_url = topic_data['videoUrl']
                
                # OPTIMIZED: O(1) quiz completion check using lookup dict
                completion_key = f"{unit_name}_{topic_name}"
                is_completed = completion_lookup.get(completion_key, False)
                
                # Determine if this topic is accessible based on refined access control logic
                if content_locked:
                    topic_is_accessible = False
                    topic_is_locked = True
                elif is_free_user:
                    if free_access_limit > 0 and free_topics_granted < free_access_limit:
                        topic_is_accessible = True
                        topic_is_locked = False
                        free_topics_granted += 1
                    else:
                        topic_is_accessible = False
                        topic_is_locked = True
                else:
                    topic_is_accessible = True
                    topic_is_locked = False

                processed_topic_data = {
                    'topicName': topic_name,
                    'videoUrl': video_url,
                    'files': files,  # Array of file objects with key, name, type
                    'hasVideo': video_url is not None,
                    'hasDocument': len(files) > 0,
                    'isAccessible': topic_is_accessible,
                    'isLocked': topic_is_locked,
                    'isCompleted': is_completed
                }
                
                topic_list.append(processed_topic_data)
            
            # Add unit with all topics (both accessible and locked)
            if topic_list:
                # Determine unit accessibility based on refined access control logic
                if content_locked:
                    # If semester access denied or trial expired, lock all units
                    unit_is_accessible = False
                    unit_is_locked = True
                elif is_free_user:
                    has_accessible_topic = any(topic['isAccessible'] for topic in topic_list)
                    unit_is_accessible = has_accessible_topic
                    unit_is_locked = not has_accessible_topic
                else:
                    unit_is_accessible = True
                    unit_is_locked = False
                
                result.append({
                    'unitName': unit_name,
                    'topics': topic_list,
                    'isAccessible': unit_is_accessible,
                    'isLocked': unit_is_locked
                })
        
        access_type = 'semester access denied' if content_locked and not trial_expired else 'free (trial expired)' if content_locked else 'free (trial active)' if is_free_user else 'premium'
        
        # Transform the result to match frontend expectations
        transformed_result = []
        for unit in result:
            transformed_unit = {
                "name": unit["unitName"],
                "isAccessible": unit.get("isAccessible", True),
                "isLocked": unit.get("isLocked", False),
                "topics": []
            }
            
            for topic in unit["topics"]:
                transformed_topic = {
                    "title": topic["topicName"],
                    "duration": "15 min",  # Default duration
                    "videoUrl": topic.get("videoUrl"),
                    "files": topic.get("files", []),
                    "hasVideo": topic.get("hasVideo", False),
                    "hasDocument": topic.get("hasDocument", False),
                    "isAccessible": topic.get("isAccessible", True),
                    "isLocked": topic.get("isLocked", False),
                    "isCompleted": topic.get("isCompleted", False)
                }
                transformed_unit["topics"].append(transformed_topic)
            
            transformed_result.append(transformed_unit)

        if is_enabled():
            try:
                capture_event(
                    current_user.mobile,  # Use mobile number as distinct_id instead of user.id
                    "subject_content_requested",
                    {
                        "course": courseName,
                        "year_semester": yearSemester,
                        "subject": subjectName,
                        "units_returned": len(transformed_result),
                        "content_locked": content_locked,
                        "free_user": is_free_user,
                        "trial_expired": trial_expired,
                        "total_topics": total_topics_count,
                        "user_id": current_user.id,  # Include user_id in properties for reference
                    },
                )
            except Exception as analytics_error:
                logger.debug("PostHog capture failed for subject content: %s", analytics_error)
        return transformed_result
        
    except Exception as e:
        print(f"ERROR: Failed to get subject content: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get subject content: {str(e)}")

@router.get('/extract-text-content')
def extract_text_content_endpoint(key: str = Query(..., description="S3 object key for file"), current_user: User = Depends(get_current_user)):
    """Extract text content from PDF or DOCX file with proper formatting including tables"""
    print(f"DEBUG: extract-text-content called with key: {key}")
    print(f"DEBUG: Current user: {current_user.mobile if current_user else 'None'}")
    
    try:
        print(f"DEBUG: Attempting to extract text content for key: {key}")
        
        # Validate the key parameter
        if not key:
            raise HTTPException(status_code=400, detail="Key parameter is required")
        
        # Check if the file exists in S3 before attempting to get it
        try:
            s3.head_object(Bucket=BUCKET, Key=key)
            print(f"DEBUG: File exists in S3: {key}")
        except Exception as s3_error:
            print(f"DEBUG: File not found in S3: {key}, Error: {str(s3_error)}")
            return JSONResponse(
                status_code=404,
                content={
                    "error": "File not found",
                    "message": f"The file '{key}' does not exist in S3",
                    "suggestions": [
                        "Check if the file path is correct",
                        "Verify the file has been uploaded to S3",
                        "Ensure the key parameter is properly URL encoded"
                    ]
                }
            )
        
        # Get the file from S3
        file_obj = s3.get_object(Bucket=BUCKET, Key=key)
        file_content = file_obj['Body'].read()
        
        print(f"DEBUG: Successfully retrieved file, size: {len(file_content)} bytes")
        
        # Extract text content based on file type
        file_extension = key.lower().split('.')[-1]
        
        if file_extension in ['docx', 'doc']:
            extracted_content = extract_docx_text_with_formatting(file_content)
        elif file_extension == 'pdf':
            extracted_content = extract_pdf_text_with_formatting(file_content)
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_extension}")
        
        print(f"DEBUG: Successfully extracted text content")
        
        # Return extracted content as JSON
        return JSONResponse(
            content={
                "filename": os.path.basename(key),
                "file_type": file_extension,
                "content": extracted_content,
                "extracted_at": datetime.datetime.now().isoformat()
            }
        )
        
    except HTTPException:
        # Re-raise HTTP exceptions as-is
        raise
    except Exception as e:
        print(f"Error in extract_text_content_endpoint: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to extract text content: {str(e)}")

def extract_docx_text_with_formatting(docx_content):
    """Extract text content from DOCX with proper formatting including tables"""
    try:
        # Create a temporary file to store the DOCX content
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            temp_docx.write(docx_content)
            temp_docx_path = temp_docx.name
        
        # Load the DOCX document
        doc = Document(temp_docx_path)
        
        # Initialize content structure
        content = {
            "paragraphs": [],
            "tables": [],
            "headers": [],
            "footers": []
        }
        
        # Extract paragraphs with formatting
        for para in doc.paragraphs:
            if para.text.strip():
                paragraph_data = {
                    "text": para.text.strip(),
                    "style": para.style.name,
                    "alignment": str(para.alignment),
                    "bold": any(run.bold for run in para.runs if run.bold),
                    "italic": any(run.italic for run in para.runs if run.italic),
                    "underline": any(run.underline for run in para.runs if run.underline),
                    "font_size": para.runs[0].font.size.pt if para.runs and para.runs[0].font.size else None
                }
                content["paragraphs"].append(paragraph_data)
        
        # Extract tables with structure
        for table_idx, table in enumerate(doc.tables):
            table_data = {
                "table_index": table_idx,
                "rows": []
            }
            
            for row_idx, row in enumerate(table.rows):
                row_data = {
                    "row_index": row_idx,
                    "cells": []
                }
                
                for cell_idx, cell in enumerate(row.cells):
                    cell_data = {
                        "cell_index": cell_idx,
                        "text": cell.text.strip(),
                        "paragraphs": []
                    }
                    
                    # Extract paragraphs from cell
                    for para in cell.paragraphs:
                        if para.text.strip():
                            cell_para = {
                                "text": para.text.strip(),
                                "style": para.style.name,
                                "bold": any(run.bold for run in para.runs if run.bold),
                                "italic": any(run.italic for run in para.runs if run.italic)
                            }
                            cell_data["paragraphs"].append(cell_para)
                    
                    row_data["cells"].append(cell_data)
                
                table_data["rows"].append(row_data)
            
            content["tables"].append(table_data)
        
        # Extract headers and footers
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        header_data = {
                            "text": para.text.strip(),
                            "style": para.style.name
                        }
                        content["headers"].append(header_data)
            
            if section.footer:
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        footer_data = {
                            "text": para.text.strip(),
                            "style": para.style.name
                        }
                        content["footers"].append(footer_data)
        
        # Clean up temporary file
        os.unlink(temp_docx_path)
        
        return content
        
    except Exception as e:
        print(f"Error extracting DOCX text with formatting: {str(e)}")
        raise e

def extract_pdf_text_with_formatting(pdf_content):
    """Extract text content from PDF with basic formatting"""
    try:
        # Create a temporary file to store the PDF content
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
            temp_pdf.write(pdf_content)
            temp_pdf_path = temp_pdf.name
        
        # Import PyPDF2 for PDF processing
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise Exception("PyPDF2 library is required for PDF text extraction")
        
        # Read PDF
        reader = PdfReader(temp_pdf_path)
        
        # Initialize content structure
        content = {
            "pages": [],
            "total_pages": len(reader.pages)
        }
        
        # Extract text from each page
        for page_idx, page in enumerate(reader.pages):
            page_data = {
                "page_number": page_idx + 1,
                "text": page.extract_text() or "",
                "page_size": {
                    "width": float(page.mediabox.width),
                    "height": float(page.mediabox.height)
                }
            }
            content["pages"].append(page_data)
        
        # Clean up temporary file
        os.unlink(temp_pdf_path)
        
        return content
        
    except Exception as e:
        print(f"Error extracting PDF text with formatting: {str(e)}")
        raise e 











@router.get('/get-notes-by-subject-topic')
def get_notes_by_subject_topic(
    subject_name: str = Query(..., description="Subject name to search for"),
    topic_name: str = Query(..., description="Topic name to search for"),
    db: Session = Depends(get_db)
):
    """Get notes content directly by subject name and topic name from the generated_notes table"""
    try:
        print(f"DEBUG: Searching for notes with subject_name: '{subject_name}', topic_name: '{topic_name}'")
        
        # 1. Try exact match first
        note = db.query(GeneratedNotes).filter(
            GeneratedNotes.subject_name == subject_name,
            GeneratedNotes.topic == topic_name
        ).first()

        # 2. Try cleaned/relaxed match if exact match fails
        if not note:
            print("DEBUG: Exact match failed, trying relaxed match...")
            from app.utils.subject_utils import clean_subject_name
            cleaned_subject = clean_subject_name(subject_name)
            
            # Search by topic name first (ignoring subject for a moment to see if it exists)
            note_by_topic = db.query(GeneratedNotes).filter(
                GeneratedNotes.topic == topic_name
            ).all()
            
            if note_by_topic:
                print(f"DEBUG: Found {len(note_by_topic)} notes for topic '{topic_name}' across all subjects.")
                for n in note_by_topic:
                    print(f"DEBUG: Topic '{topic_name}' exists in subject: '{n.subject_name}'")
            
            # Use a more relaxed subject search
            subject_parts = subject_name.split(':')
            subject_core = subject_parts[-1].strip() if len(subject_parts) > 1 else subject_name
            # Remove the long dash if present
            subject_core = subject_core.replace('–', '-').split('-')[0].strip()
            
            print(f"DEBUG: Relaxed subject core: '{subject_core}'")

            note = db.query(GeneratedNotes).filter(
                or_(
                    func.lower(GeneratedNotes.subject_name).contains(func.lower(subject_core)),
                    func.lower(GeneratedNotes.topic) == func.lower(topic_name)
                )
            ).filter(func.lower(GeneratedNotes.topic) == func.lower(topic_name)).first()
        
        if note:
            # Update last accessed time
            note.updated_at = datetime.datetime.now()
            db.commit()
            
            print(f"DEBUG: Found notes for subject: {subject_name}, topic: {topic_name}")
            return {
                "found": True,
                "notes_content": note.notes_content,
                "filename": note.document_name,
                "subject": note.subject_name,
                "unit": note.unit_name,
                "topic": note.topic,
                "generation_time": note.created_at.isoformat() if note.created_at else None,
                "access_count": 1  # Default since we don't have this field
            }
        else:
            print(f"DEBUG: No notes found for subject: {subject_name}, topic: {topic_name}")
            return {
                "found": False,
                "message": "No notes available for this subject and topic"
            }
            
    except Exception as e:
        print(f"Error getting notes by subject and topic: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get notes: {str(e)}")

@router.get('/get-notes-by-document-id')
def get_notes_by_document_id(
    document_id: str = Query(..., description="Document ID to search for"),
    db: Session = Depends(get_db)
):
    """Get notes content by document ID from the generated_notes table"""
    try:
        # Search for notes in the generated_notes table
        note = db.query(GeneratedNotes).filter(
            GeneratedNotes.document_id == document_id
        ).first()
        
        if note:
            # Update last accessed time
            note.updated_at = datetime.datetime.now()
            db.commit()
            
            return {
                "found": True,
                "notes_content": note.notes_content,
                "filename": note.document_name,
                "subject": note.subject_name,
                "unit": note.unit_name,
                "topic": note.topic,
                "generation_time": note.created_at.isoformat() if note.created_at else None,
                "access_count": 1  # Default since we don't have this field
            }
        else:
            return {
                "found": False,
                "message": "No notes available for this document"
            }
            
    except Exception as e:
        print(f"Error getting notes by document ID: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get notes: {str(e)}")

@router.get('/list-subjects')
def list_subjects(
    courseName: str,
    yearSemester: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """List available subjects for a given course and yearSemester (S3 prefixes)."""
    try:
        # Access control (semester-based)
        subscription = db.query(Subscription).filter(
            Subscription.user_id == current_user.id,
            Subscription.status == "active",
            Subscription.end_date > datetime.datetime.utcnow()
        ).first()

        access_info = get_content_access_info(current_user, yearSemester, subscription)
        if access_info.get("locked_by_semester") or access_info.get("locked_by_trial"):
            raise HTTPException(status_code=403, detail="Access denied for this semester content")

        # List S3 prefixes for subjects under course/yearSemester
        prefix = f"{courseName}/{yearSemester}/"
        response = s3.list_objects_v2(Bucket=BUCKET, Prefix=prefix, Delimiter='/')
        subject_prefixes = response.get('CommonPrefixes', [])

        subjects = []
        for p in subject_prefixes:
            pref = p.get('Prefix') or ''
            parts = pref.rstrip('/').split('/')
            if len(parts) >= 3:
                subjects.append(parts[-1])

        subjects = sorted(subjects)
        
        # Don't filter subjects - show all subjects so users can get free trial access
        # Subject-specific access control is handled in the individual subject content endpoint
        print(f"DEBUG: Showing all {len(subjects)} subjects for free trial access")
        
        return {"subjects": subjects}
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error listing subjects: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to list subjects")


@router.get('/progress')
def get_subject_progress(
    courseName: str,
    yearSemester: str,
    subjectName: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Return progress for a subject: total topics, completed topics, and percentage."""
    try:
        # Access control
        subscription = db.query(Subscription).filter(
            Subscription.user_id == current_user.id,
            Subscription.status == "active",
            Subscription.end_date > datetime.datetime.utcnow()
        ).first()
        access_info = get_content_access_info(current_user, yearSemester, subscription)
        if access_info.get("locked_by_semester") or access_info.get("locked_by_trial"):
            raise HTTPException(status_code=403, detail="Access denied for this semester content")

        # Get content structure for this subject (cached internally)
        # DISABLED: No longer using S3 - use curriculum API instead
        content = []

        # Count total topics from structure
        total_topics = 0
        for unit in content:
            topics = unit.get('topics', [])
            total_topics += len(topics)

        # Completed topics from Quiz table (distinct topics with XP)
        from app.utils.subject_utils import clean_subject_name
        cleaned_subject = clean_subject_name(subjectName)
        completed_topics = (
            db.query(Quiz.topic)
            .filter(
                Quiz.student_id == current_user.id,
                Quiz.subject == cleaned_subject,
                Quiz.xp_topic.isnot(None),
            )
            .distinct()
            .count()
        )

        progress = min(100, int((completed_topics / total_topics) * 100)) if total_topics > 0 else 0

        return {
            "subject": subjectName,
            "totalTopics": total_topics,
            "completed": int(completed_topics or 0),
            "progress": progress,
        }
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error computing subject progress: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to compute subject progress")

# ===== Accessible semesters helpers and endpoint =====

def get_accessible_semesters(year, semester):
    """
    Compute accessible semesters based on user's current year and semester.
    Returns list like ["1-1", "1-2", "2-1"] up to current semester.
    """
    try:
        year = int(year or 1)
        semester = int(semester or 1)
    except Exception:
        year, semester = 1, 1
    accessible = []
    for y in range(1, year + 1):
        if y < year:
            accessible.extend([f"{y}-1", f"{y}-2"])
        else:
            for s in range(1, semester + 1):
                accessible.append(f"{y}-{s}")
    return accessible

async def get_user_accessible_semesters(current_user):
    """
    Get accessible semesters for the current user based on their year/semester.
    """
    y = getattr(current_user, "year", None)
    s = getattr(current_user, "semester", None)
    if not y or not s:
        return []
    return get_accessible_semesters(y, s)

def is_semester_accessible(current_user, year_semester: str) -> bool:
    """
    Check if a given year-semester string is accessible to the user.
    Future semesters should be blocked.
    """
    try:
        y = int(getattr(current_user, "year", 1) or 1)
        s = int(getattr(current_user, "semester", 1) or 1)
    except Exception:
        return True
    accessible = get_accessible_semesters(y, s)
    return year_semester in accessible

@router.get("/accessible-semesters")
async def get_accessible_semesters_endpoint(
    current_user: User = Depends(get_current_user),
):
    """
    List of semesters accessible to the current user (current + completed).
    """
    return await get_user_accessible_semesters(current_user)

