from fastapi import APIRouter, HTTPException, Depends, Body
from app.schemas.quiz import QuizRequest, QuizResponse, ScoreRequest, ScoreResponse
from app.api.auth import get_current_user
from app.models.user import User
from app.models.quiz import Quiz, QuizAttemptedQuestion
from sqlalchemy.orm import Session
from app.db.session import get_db
from datetime import datetime, date
import logging
import boto3
import os
from PyPDF2 import PdfReader
from docx import Document
import tempfile
from pydantic import BaseModel
from app.utils.google_quiz import generate_quiz_questions, generate_quiz_questions_from_text
from app.services.quiz_analysis_service import quiz_analysis_service
from app.utils.subject_utils import clean_subject_name, extract_subject_code
from sqlalchemy import desc, func
import fitz  # PyMuPDF for better PDF text extraction
import io

from app.services.analytics.posthog_client import capture_event, is_enabled

logger = logging.getLogger(__name__)

router = APIRouter()
# Get environment variables
BUCKET = os.getenv('S3_BUCKET')
AWS_ACCESS_KEY_ID = os.getenv('AWS_ACCESS_KEY_ID')
AWS_SECRET_ACCESS_KEY = os.getenv('AWS_SECRET_ACCESS_KEY')
AWS_REGION = os.getenv('AWS_REGION', 'ap-south-1')

# Validate required environment variables
if not BUCKET:
    raise ValueError("S3_BUCKET environment variable is required")
if not AWS_ACCESS_KEY_ID:
    raise ValueError("AWS_ACCESS_KEY_ID environment variable is required")
if not AWS_SECRET_ACCESS_KEY:
    raise ValueError("AWS_SECRET_ACCESS_KEY environment variable is required")

class FileKeyRequest(BaseModel):
    file_key: str

def extract_text_from_pdf_pymupdf(pdf_path):
    """Extract text from PDF using PyMuPDF (more reliable than PyPDF2)"""
    try:
        doc = fitz.open(pdf_path)
        text = ""
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += page.get_text()
        doc.close()
        return text.strip()
    except Exception as e:
        logger.warning(f"PyMuPDF extraction failed: {str(e)}")
        return None

def extract_text_from_pdf_pypdf2(pdf_path):
    """Extract text from PDF using PyPDF2 (fallback method)"""
    try:
        with open(pdf_path, 'rb') as f:
            reader = PdfReader(f)
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()
    except Exception as e:
        logger.warning(f"PyPDF2 extraction failed: {str(e)}")
        return None

def extract_text_from_pdf(pdf_path):
    """Extract text from PDF using multiple methods with fallbacks"""
    logger.info(f"Attempting to extract text from PDF: {pdf_path}")
    
    # Validate PDF file first
    try:
        with open(pdf_path, 'rb') as f:
            header = f.read(4)
            if header != b'%PDF':
                logger.error(f"File doesn't appear to be a valid PDF. Header: {header}")
                return None
    except Exception as e:
        logger.error(f"Error reading PDF file header: {str(e)}")
        return None
    
    # Try PyMuPDF first (more reliable)
    text = extract_text_from_pdf_pymupdf(pdf_path)
    if text and len(text.strip()) > 10:  # Ensure we got meaningful text
        logger.info(f"Successfully extracted {len(text)} characters using PyMuPDF")
        return text
    
    # Fallback to PyPDF2
    text = extract_text_from_pdf_pypdf2(pdf_path)
    if text and len(text.strip()) > 10:
        logger.info(f"Successfully extracted {len(text)} characters using PyPDF2")
        return text
    
    # If both methods fail, try to get at least some text
    logger.warning("Both PDF extraction methods failed, attempting minimal extraction")
    try:
        with open(pdf_path, 'rb') as f:
            reader = PdfReader(f)
            if len(reader.pages) > 0:
                # Try to get metadata or basic info
                info = reader.metadata
                logger.info(f"PDF metadata: {info}")
                
                # Try to extract text from first page only
                first_page = reader.pages[0]
                text = first_page.extract_text() or ""
                if text.strip():
                    logger.info(f"Extracted {len(text)} characters from first page only")
                    return text
                
                # If still no text, check if it might be a scanned document
                logger.warning("No text extracted - this might be a scanned document or image-based PDF")
                
    except Exception as e:
        logger.error(f"Minimal PDF extraction also failed: {str(e)}")
    
    return None
 
@router.post('/generate', response_model=QuizResponse)
def generate_quiz(
    request: QuizRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    try:
        logger.info(f"Received quiz generation request: year={request.year}, semester={request.semester}")
        # Validate year and semester
        if not isinstance(request.year, int) or request.year < 1 or request.year > 4:
            raise HTTPException(status_code=422, detail=f"Invalid year: {request.year}. Must be 1-4")
        if not isinstance(request.semester, int) or request.semester < 1 or request.semester > 2:
            raise HTTPException(status_code=422, detail=f"Invalid semester: {request.semester}. Must be 1-2")
        
        # Check if quiz already exists for this user and topic (for pre-assessment)
        topic = getattr(request, 'topic', None) or "Pre-Assessment"
        existing_quiz = db.query(Quiz).filter(
            Quiz.student_id == current_user.id,
            Quiz.topic == topic
        ).first()
        
        if existing_quiz:
            logger.info(f"Quiz already exists for user {current_user.id}, topic: {topic}")
            # Get the attempted questions for this quiz
            attempted_questions = db.query(QuizAttemptedQuestion).filter(
                QuizAttemptedQuestion.quiz_id == existing_quiz.id
            ).all()
            
            # Convert to quiz format
            questions = []
            for q in attempted_questions:
                questions.append({
                    "question": q.question_text,
                    "options": ["Option A", "Option B", "Option C", "Option D"],  # Placeholder options
                    "answer": 0  # Placeholder answer
                })
            if is_enabled():
                try:
                    capture_event(
                        str(current_user.id),
                        "quiz_existing_returned",
                        {
                            "quiz_id": existing_quiz.id,
                            "topic": topic,
                            "subject": existing_quiz.subject,
                            "score": existing_quiz.score,
                            "question_count": len(questions),
                        },
                    )
                except Exception as analytics_error:
                    logger.debug("PostHog capture failed for existing quiz: %s", analytics_error)
            return QuizResponse(
                questions=questions,
                existing_quiz=True,
                previous_score=existing_quiz.score,
                quiz_id=existing_quiz.id,
                message=f"You have already taken this quiz. Your previous score was {existing_quiz.score}/5.",
                topic=topic
            )
        
        # Use question bank for prelogin assessment - NO OpenAI API key required
        questions = generate_quiz_questions(request.year, request.semester)
        logger.info(f"Generated quiz questions from question bank (no OpenAI): {questions}")
        if is_enabled():
            try:
                capture_event(
                    str(current_user.id),
                    "quiz_generated",
                    {
                        "year": request.year,
                        "semester": request.semester,
                        "subject": request.subject,
                        "unit": request.unit,
                        "topic": request.topic,
                        "question_count": len(questions),
                    },
                )
            except Exception as analytics_error:
                logger.debug("PostHog capture failed for quiz generation: %s", analytics_error)
        return QuizResponse(questions=questions)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generating quiz: {str(e)}")
        if is_enabled():
            try:
                capture_event(
                    str(current_user.id),
                    "quiz_generation_failed",
                    {
                        "year": request.year,
                        "semester": request.semester,
                        "subject": request.subject,
                        "reason": str(e),
                    },
                )
            except Exception as analytics_error:
                logger.debug("PostHog capture failed for quiz generation error: %s", analytics_error)
        raise HTTPException(status_code=500, detail=str(e))
 
@router.get('/scores')
def get_quiz_scores(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Get all quiz scores for the current user"""
    try:
        logger.info(f"Fetching quiz scores for user {current_user.id}")
        
        # Get all quiz records for this user
        quiz_records = db.query(Quiz).filter_by(student_id=current_user.id).all()
        
        scores = []
        for quiz in quiz_records:
            scores.append({
                "id": quiz.id,
                "year": quiz.year,
                "semester": quiz.semester,
                "subject": quiz.subject,
                "unit": quiz.unit,
                "topic": quiz.topic,
                "score": quiz.score,
                "strengths": quiz.strengths,
                "weakness": quiz.weakness,
                "areas_to_improve": quiz.areas_to_improve,
                "created_at": quiz.created_at.isoformat() if quiz.created_at else None
            })
        
        logger.info(f"Found {len(scores)} quiz records for user {current_user.id}")
        return {"scores": scores}
        
    except Exception as e:
        logger.error(f"Error fetching quiz scores: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post('/scores')
async def create_quiz_score(
    request: ScoreRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Create a new quiz score record - alternative endpoint to /score"""
    try:
        logger.info(f"Creating quiz score via /scores endpoint for user {current_user.id}")
        
        # Calculate score
        score = 0
        results = []
        for q, user_ans in zip(request.questions, request.user_answers):
            correct = (user_ans == q.get('answer'))
            results.append(correct)
            if correct:
                score += 1
        
        # Delete previous quiz for this user/topic
        existing_quiz = db.query(Quiz).filter_by(
            student_id=current_user.id,
            subject=request.subject or "General",
            unit=request.unit or "General",
            topic=request.topic or "General"
        ).first()
        
        if existing_quiz:
            db.query(QuizAttemptedQuestion).filter_by(quiz_id=existing_quiz.id).delete()
            db.delete(existing_quiz)
            db.commit()
            logger.info(f"Deleted existing quiz {existing_quiz.id}")
        
        # Generate quiz analysis using Gemini API
        analysis = None
        try:
            logger.info(f"Generating quiz analysis for user {current_user.id}, score: {score}")
            analysis = await quiz_analysis_service.analyze_quiz_performance(
                questions=request.questions,
                user_answers=request.user_answers,
                score=score,
                subject=request.subject or "General",
                topic=request.topic or "General"
            )
            logger.info(f"Quiz analysis generated successfully: {analysis}")
        except Exception as e:
            logger.error(f"Failed to generate quiz analysis: {str(e)}")
            # Continue without analysis if it fails
            analysis = {
                'strengths': 'Analysis unavailable',
                'weakness': 'Analysis unavailable', 
                'areas_to_improve': 'Analysis unavailable'
            }

        # Create new quiz record with analysis
        quiz_record = Quiz(
            student_id=current_user.id,
            year=request.year,
            semester=request.semester,
            subject=clean_subject_name(request.subject) or "General",
            unit=request.unit or "General",
            topic=request.topic or "General",
            score=score,
            xp_topic=0,  # Default XP for this endpoint
            strengths=analysis.get('strengths') if analysis else None,
            weakness=analysis.get('weakness') if analysis else None,
            areas_to_improve=analysis.get('areas_to_improve') if analysis else None
        )
        db.add(quiz_record)
        db.commit()
        db.refresh(quiz_record)
        
        # Update user's quiz_score if higher
        if current_user.quiz_score is None or score > current_user.quiz_score:
            current_user.quiz_score = score
            db.add(current_user)
            db.commit()
        
        # Store attempted questions
        for q in request.questions:
            question_text = q.get('question')
            if question_text:
                attempted_q = QuizAttemptedQuestion(
                    quiz_id=quiz_record.id,
                    question_text=question_text
                )
                db.add(attempted_q)
        db.commit()
        
        logger.info(f"Successfully created quiz score: {score} for user {current_user.id}")
        event_payload = {
            "score": score,
            "total": len(request.questions),
            "results": results,
            "xp_earned": xp_earned,
            "xp_breakdown": xp_breakdown
        }
        if is_enabled():
            try:
                capture_event(
                    str(current_user.id),
                    "quiz_completed",
                    {
                        "score": score,
                        "total_questions": len(request.questions),
                        "subject": request.subject,
                        "unit": request.unit,
                        "topic": request.topic,
                        "xp_earned": xp_earned,
                        "is_onboarding": is_onboarding_quiz,
                    },
                )
            except Exception as analytics_error:
                logger.debug("PostHog capture failed for quiz completion: %s", analytics_error)
        return event_payload

    except Exception as e:
        logger.error(f"Error creating quiz score: {str(e)}")
        if is_enabled():
            try:
                capture_event(
                    str(current_user.id),
                    "quiz_score_failed",
                    {
                        "subject": request.subject,
                        "unit": request.unit,
                        "topic": request.topic,
                        "reason": str(e),
                    },
                )
            except Exception as analytics_error:
                logger.debug("PostHog capture failed for quiz score error: %s", analytics_error)
        raise HTTPException(status_code=500, detail=str(e))

@router.post('/score', response_model=ScoreResponse)
async def calculate_score(
    request: ScoreRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    try:
        logger.info("=== Starting Score Calculation ===")
        logger.info(f"Received questions: {request.questions}")
        logger.info(f"Received user answers: {request.user_answers}")
       
        score = 0
        results = []
        for i, (q, user_ans) in enumerate(zip(request.questions, request.user_answers)):
            logger.info(f"\nQuestion {i + 1}:")
            logger.info(f"Question text: {q.get('question')}")
            logger.info(f"Options: {q.get('options')}")
            logger.info(f"Correct answer index: {q.get('answer')}")
            logger.info(f"User answer index: {user_ans}")
           
            correct = (user_ans == q.get('answer'))
            logger.info(f"Is correct: {correct}")
           
            results.append(correct)
            if correct:
                score += 1
                logger.info(f"Current score: {score}")
       
        logger.info(f"\nFinal score: {score} out of {len(request.questions)}")
        logger.info(f"Results array: {results}")
        logger.info(f"Quiz details - Subject: {request.subject}, Unit: {request.unit}, Topic: {request.topic}")
       
        # Delete previous quiz (and its questions) for this user/topic
        # First find the existing quiz - be more specific for onboarding quizzes
        existing_quiz = None
        
        # For onboarding quizzes, look for exact matches
        if request.subject == "Pre-Assessment" and request.unit == "Onboarding" and request.topic == "Onboarding":
            logger.info(f"Looking for existing onboarding quiz for user {current_user.id}")
            existing_quiz = db.query(Quiz).filter(
                Quiz.student_id == current_user.id,
                Quiz.subject == "Pre-Assessment",
                Quiz.unit == "Onboarding",
                Quiz.topic == "Onboarding"
            ).first()
            logger.info(f"Found existing onboarding quiz: {existing_quiz.id if existing_quiz else 'None'}")
        else:
            # For other quizzes, use the original logic
            logger.info(f"Looking for existing quiz for user {current_user.id} with subject={request.subject}, unit={request.unit}, topic={request.topic}")
            existing_quiz = db.query(Quiz).filter_by(
                student_id=current_user.id,
                subject=request.subject or "General",
                unit=request.unit or "General",
                topic=request.topic or "General"
            ).first()
            logger.info(f"Found existing quiz: {existing_quiz.id if existing_quiz else 'None'}")
        
        if existing_quiz:
            # Delete the attempted questions first (to avoid foreign key constraint)
            db.query(QuizAttemptedQuestion).filter_by(quiz_id=existing_quiz.id).delete()
            # Then delete the quiz record
            db.delete(existing_quiz)
            db.commit()
            logger.info(f"Deleted existing quiz {existing_quiz.id} and its questions for user {current_user.id}")
 
        # Calculate XP for onboarding quiz: XP = score * 2
        xp_earned = 0
        xp_breakdown = {}
        
        # Check for onboarding quiz using both subject and topic combinations
        is_onboarding_quiz = (
            (request.subject == "Pre-Assessment" and request.unit == "Onboarding" and request.topic == "Onboarding") or
            request.topic == "Pre-Assessment"
        )
        
        if is_onboarding_quiz:
            xp_earned = score * 2
            xp_breakdown = {
                "quiz_score": score,
                "multiplier": 2,
                "total": xp_earned
            }
            current_user.total_xp += xp_earned
            db.add(current_user)
            logger.info(f"Awarded {xp_earned} XP for onboarding quiz completion (score: {score})")

        # Calculate XP for course content quiz: XP = video(10) + notes(5) + quiz_score(score*1)
        # Only for non-onboarding quizzes
        if not is_onboarding_quiz:
            # When quiz is completed, video and notes are automatically considered completed
            video_xp = 10  # Fixed XP for video completion
            notes_xp = 5   # Fixed XP for notes completion
            quiz_xp = score * 1  # score * 1 for quiz completion
            
            xp_earned = video_xp + notes_xp + quiz_xp
            xp_breakdown = {
                "video_xp": video_xp,
                "notes_xp": notes_xp,
                "quiz_xp": quiz_xp,
                "total": xp_earned
            }
            
            if xp_earned > 0:
                current_user.total_xp += xp_earned
                db.add(current_user)
                logger.info(f"Awarded {xp_earned} XP for course content quiz completion (video: {video_xp}, notes: {notes_xp}, quiz: {quiz_xp})")

        # Generate quiz analysis using Gemini API
        analysis = None
        try:
            logger.info(f"Generating quiz analysis for user {current_user.id}, score: {score}")
            analysis = await quiz_analysis_service.analyze_quiz_performance(
                questions=request.questions,
                user_answers=request.user_answers,
                score=score,
                subject=request.subject or "General",
                topic=request.topic or "General"
            )
            logger.info(f"Quiz analysis generated successfully: {analysis}")
        except Exception as e:
            logger.error(f"Failed to generate quiz analysis: {str(e)}")
            # Continue without analysis if it fails
            analysis = {
                'strengths': 'Analysis unavailable',
                'weakness': 'Analysis unavailable', 
                'areas_to_improve': 'Analysis unavailable'
            }

        # Save quiz record with analysis
        quiz_record = Quiz(
            student_id=current_user.id,
            year=request.year,
            semester=request.semester,
            subject=clean_subject_name(request.subject) or "General",
            unit=request.unit or "General",
            topic=request.topic or "General",
            score=score,
            xp_topic=xp_earned,
            strengths=analysis.get('strengths') if analysis else None,
            weakness=analysis.get('weakness') if analysis else None,
            areas_to_improve=analysis.get('areas_to_improve') if analysis else None
        )
        db.add(quiz_record)
        db.commit()
        db.refresh(quiz_record)
       
        # Always update user's year and semester
        current_user.year = request.year
        current_user.semester = request.semester
        db.add(current_user)
        # Only update quiz_score if this is their first quiz or if they got a higher score
        if current_user.quiz_score is None or score > current_user.quiz_score:
            current_user.quiz_score = score
            db.add(current_user)
       
        # Store new attempted questions for this quiz
        for q in request.questions:
            question_text = q.get('question')
            if question_text:
                attempted_q = QuizAttemptedQuestion(
                    quiz_id=quiz_record.id,
                    question_text=question_text
                )
                db.add(attempted_q)
        
        # Single commit for all changes with proper error handling
        try:
            db.commit()
            db.refresh(quiz_record)
        except Exception as e:
            logger.error(f"Failed to commit quiz data: {str(e)}")
            db.rollback()
            raise HTTPException(status_code=500, detail="Failed to save quiz data")
 
        logger.info(f"Saved quiz score to PostgreSQL quiz table for student {current_user.id} with score {score}")
        logger.info(f"Updated user's quiz_score to {current_user.quiz_score}")
        
        # Update study plan task quiz_completed status
        if not is_onboarding_quiz:
            try:
                from app.models.study_plan import StudyPlanTask
                # Note: clean_subject_name is imported at module level to avoid shadowing inside this function
                from sqlalchemy import and_, or_
                
                # Get cleaned subject name and code for matching
                subject_code = extract_subject_code(request.subject)
                cleaned_subject = clean_subject_name(request.subject)
                
                logger.info(f"Searching for study plan task with subject='{request.subject}', cleaned='{cleaned_subject}', code='{subject_code}', topic='{request.topic}'")
                
                # Build matching conditions - study plan uses full names with codes
                subject_match_conditions = []
                if subject_code:
                    subject_match_conditions.append(StudyPlanTask.subject.contains(subject_code))
                if cleaned_subject:
                    subject_match_conditions.append(StudyPlanTask.subject.contains(cleaned_subject))
                if request.subject:
                    subject_match_conditions.append(StudyPlanTask.subject.contains(request.subject))
                
                topic_match_conditions = [
                    StudyPlanTask.title.contains(request.topic),
                    StudyPlanTask.notes.contains(f"Topic: {request.topic}"),
                    StudyPlanTask.notes.ilike(f"%{request.topic}%")
                ]
                
                # Try to find the corresponding study plan task
                study_plan_task = db.query(StudyPlanTask).filter(
                    and_(
                        StudyPlanTask.user_id == current_user.id,
                        or_(*subject_match_conditions) if subject_match_conditions else True,
                        or_(*topic_match_conditions)
                    )
                ).first()
                
                if study_plan_task:
                    logger.info(f"Found study plan task: {study_plan_task.task_id}")
                    
                    # Mark quiz as completed - this also marks the entire task as completed
                    study_plan_task.quiz_completed = True
                    study_plan_task.completed = True
                    study_plan_task.completed_at = datetime.utcnow()
                    study_plan_task.updated_at = datetime.utcnow()
                    
                    db.commit()
                    logger.info(f"✅ Updated study plan task {study_plan_task.task_id}: quiz_completed=True, completed=True")
                else:
                    logger.warning(f"No study plan task found for topic: {request.topic}, subject: {request.subject}")
            except Exception as e:
                logger.error(f"Failed to update study plan task: {str(e)}")
                # Don't fail the whole quiz scoring if study plan update fails
                import traceback
                traceback.print_exc()
       
        return ScoreResponse(
            score=score, 
            total=len(request.questions), 
            results=results,
            xp_earned=xp_earned,
            xp_breakdown=xp_breakdown
        )
    except Exception as e:
        logger.error(f"Exception in /score: {str(e)}")
        import traceback; traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))
 
@router.post('/generate-from-bank-public', response_model=QuizResponse)
def generate_quiz_from_bank_public(request: QuizRequest):
    """
    Generate quiz questions from the question bank for public/prelogin assessment.
    This endpoint does NOT require authentication - for onboarding use.
    """
    try:
        logger.info(f"Received public quiz generation request: year={request.year}, semester={request.semester}")
        
        from app.utils.question_bank import get_random_questions
        questions = get_random_questions(request.year, request.semester, 5)
        logger.info(f"Generated public quiz questions from question bank: {questions}")
        return QuizResponse(questions=questions)
    except Exception as e:
        logger.error(f"Error generating public quiz from question bank: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post('/generate-from-bank', response_model=QuizResponse)
def generate_quiz_from_bank(
    request: QuizRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Generate quiz questions from the question bank for prelogin assessment.
    This endpoint uses ONLY the question bank - NO OpenAI API key required.
    """
    try:
        # Check if quiz already exists for this user and topic
        topic = getattr(request, 'topic', None) or "Pre-Assessment"
        existing_quiz = db.query(Quiz).filter(
            Quiz.student_id == current_user.id,
            Quiz.topic == topic
        ).first()
        
        if existing_quiz:
            logger.info(f"Quiz already exists for user {current_user.id}, topic: {topic}")
            # Get the attempted questions for this quiz
            attempted_questions = db.query(QuizAttemptedQuestion).filter(
                QuizAttemptedQuestion.quiz_id == existing_quiz.id
            ).all()
            
            # Convert to quiz format
            questions = []
            for q in attempted_questions:
                questions.append({
                    "question": q.question_text,
                    "options": ["Option A", "Option B", "Option C", "Option D"],  # Placeholder options
                    "answer": 0  # Placeholder answer
                })
            
            return QuizResponse(
                questions=questions,
                existing_quiz=True,
                previous_score=existing_quiz.score,
                quiz_id=existing_quiz.id,
                message=f"You have already taken this quiz for {topic}. Your previous score was {existing_quiz.score}/5. You can only take each topic quiz once.",
                topic=topic
            )
        
        from app.utils.question_bank import get_random_questions
        questions = get_random_questions(request.year, request.semester, 5)
        logger.info(f"Generated quiz questions from question bank (no OpenAI): {questions}")
        return QuizResponse(questions=questions)
    except Exception as e:
        logger.error(f"Error generating quiz from question bank: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
 
@router.post('/generate-from-file', response_model=QuizResponse)
async def generate_quiz_from_file(
    request: FileKeyRequest,
    current_user: User = Depends(get_current_user),  # 🔐 ADD AUTHENTICATION
    db: Session = Depends(get_db)
):
    try:
        file_key = request.file_key
        logger.info(f"Received file_key: {file_key} from user {current_user.mobile}")
        
        # Extract subject and topic information from file key
        # Path structure: bpharma/pci/2_1/Pharmaceutical Organic Chemistry –II/Benzene and its Derivatives/Benzene/Syllabus_B_Pharm.pdf
        # OR: bpharmacy/1-2/BS203: Biochemistry/Unit 1/Biological Oxidation/Biological_Oxidation.docx
        path_parts = file_key.split('/')
        
        # Detect path structure: if path_parts[2] looks like year_semester (e.g., "2_1", "1-2"), use new structure
        year_semester_pattern = r'^\d+[_-]\d+$'
        import re
        
        if len(path_parts) >= 6 and re.match(year_semester_pattern, path_parts[2]):
            # New structure: bpharma/pci/2_1/Subject/Unit/Topic/file.pdf
            # path_parts[0] = "bpharma" (course)
            # path_parts[1] = "pci" (university)
            # path_parts[2] = "2_1" (year_semester - SKIP THIS)
            # path_parts[3] = "Pharmaceutical Organic Chemistry –II" (SUBJECT)
            # path_parts[4] = "Benzene and its Derivatives" (UNIT)
            # path_parts[5] = "Benzene" (TOPIC)
            subject_part = path_parts[3]  # Subject name
            unit_part = path_parts[4]     # Unit name
            topic_part = path_parts[5]     # Topic name
        elif len(path_parts) >= 5:
            # Old structure: bpharmacy/1-2/BS203: Biochemistry/Unit 1/Biological Oxidation/file.docx
            subject_part = path_parts[2]  # BS203: Biochemistry
            unit_part = path_parts[3]     # Unit 1
            topic_part = path_parts[4]   # Biological Oxidation
        else:
            logger.warning(f"Could not extract subject/topic from file key (insufficient path parts): {file_key}")
            subject_part = "Unknown"
            unit_part = "Unknown"
            topic_part = "Unknown"
        
        # Clean up subject name (remove course code if present)
        subject = clean_subject_name(subject_part)
        unit = unit_part
        topic = topic_part
        
        logger.info(f"Extracted from path '{file_key}' - Subject: {subject}, Unit: {unit}, Topic: {topic}")
        
        # Check if quiz already exists for this user, subject, unit, and topic
        existing_quiz = db.query(Quiz).filter(
            Quiz.student_id == current_user.id,
            Quiz.subject == subject,
            Quiz.unit == unit,
            Quiz.topic == topic
        ).first()
        
        if existing_quiz:
            logger.info(f"Quiz already exists for user {current_user.id}, subject: {subject}, unit: {unit}, topic: {topic}")
            # Get the attempted questions for this quiz
            attempted_questions = db.query(QuizAttemptedQuestion).filter(
                QuizAttemptedQuestion.quiz_id == existing_quiz.id
            ).all()
            
            # Convert to quiz format
            questions = []
            for q in attempted_questions:
                questions.append({
                    "question": q.question_text,
                    "options": ["Option A", "Option B", "Option C", "Option D"],  # Placeholder options
                    "answer": 0  # Placeholder answer
                })
            
            return QuizResponse(
                questions=questions,
                existing_quiz=True,
                previous_score=existing_quiz.score,
                quiz_id=existing_quiz.id,
                message=f"You have already taken this quiz for {topic}. Your previous score was {existing_quiz.score}/5. You can only take each topic quiz once.",
                subject=subject,
                unit=unit,
                topic=topic
            )
 
        # Fetch S3 credentials from environment at runtime
        BUCKET = os.getenv('S3_BUCKET')
        AWS_ACCESS_KEY_ID = os.getenv('AWS_ACCESS_KEY_ID')
        AWS_SECRET_ACCESS_KEY = os.getenv('AWS_SECRET_ACCESS_KEY')
        AWS_REGION = os.getenv('AWS_REGION', 'ap-south-1')
        
        # Validate S3 configuration
        if not BUCKET:
            logger.error("S3_BUCKET environment variable is not set")
            raise HTTPException(status_code=500, detail="S3 configuration error: BUCKET not set")
        if not AWS_ACCESS_KEY_ID:
            logger.error("AWS_ACCESS_KEY_ID environment variable is not set")
            raise HTTPException(status_code=500, detail="S3 configuration error: ACCESS_KEY_ID not set")
        if not AWS_SECRET_ACCESS_KEY:
            logger.error("AWS_SECRET_ACCESS_KEY environment variable is not set")
            raise HTTPException(status_code=500, detail="S3 configuration error: SECRET_ACCESS_KEY not set")
        
        logger.info(f"S3 config - BUCKET: {BUCKET}, REGION: {AWS_REGION}, ACCESS_KEY_ID: {'***' if AWS_ACCESS_KEY_ID else 'None'}")
 
        # Download file from S3 with error handling
        try:
            s3 = boto3.client(
                's3',
                aws_access_key_id=AWS_ACCESS_KEY_ID,
                aws_secret_access_key=AWS_SECRET_ACCESS_KEY,
                region_name=AWS_REGION
            )
            logger.info("S3 client created successfully")
            
            with tempfile.NamedTemporaryFile(delete=False) as tmp:
                logger.info(f"Attempting to download file {file_key} from bucket {BUCKET}")
                s3.download_fileobj(BUCKET, file_key, tmp)
                tmp_path = tmp.name
            logger.info(f"File downloaded to {tmp_path}")
            
        except Exception as s3_error:
            logger.error(f"S3 download failed: {str(s3_error)}")
            raise HTTPException(status_code=500, detail=f"Failed to download file from S3: {str(s3_error)}")
 
        # Extract text with better error handling
        text = ''
        try:
            if file_key.lower().endswith('.pdf'):
                logger.info(f"Processing PDF file: {file_key}")
                logger.info(f"PDF file size: {os.path.getsize(tmp_path)} bytes")
                text = extract_text_from_pdf(tmp_path)
                if text:
                    logger.info(f"PDF text extraction successful: {len(text)} characters")
                else:
                    logger.error("PDF text extraction failed - no text extracted")
            elif file_key.lower().endswith('.docx'):
                try:
                    doc = Document(tmp_path)
                    logger.info(f"Processing DOCX file with {len(doc.paragraphs)} paragraphs")
                   
                    # Extract text from paragraphs
                    for para in doc.paragraphs:
                        if para.text.strip():  # Only add non-empty paragraphs
                            text += para.text.strip() + '\n'
                   
                    # Extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():  # Only add non-empty cells
                                    text += cell.text.strip() + ' '
                            text += '\n'  # New line after each row
                   
                    # Extract text from headers and footers
                    for section in doc.sections:
                        if section.header:
                            for para in section.header.paragraphs:
                                if para.text.strip():
                                    text += para.text.strip() + '\n'
                        if section.footer:
                            for para in section.footer.paragraphs:
                                if para.text.strip():
                                    text += para.text.strip() + '\n'
                   
                    logger.info(f"Extracted {len(text)} characters from DOCX file")
                   
                    if not text.strip():
                        logger.warning("No text content extracted from DOCX file")
                        raise HTTPException(status_code=400, detail='No text content found in DOCX file')
                       
                except Exception as docx_error:
                    logger.error(f"Error processing DOCX file: {str(docx_error)}")
                    raise HTTPException(status_code=500, detail=f'Error processing DOCX file: {str(docx_error)}')
            else:
                os.unlink(tmp_path)
                raise HTTPException(status_code=400, detail='Unsupported file type')
        except Exception as text_error:
            logger.error(f"Text extraction failed: {str(text_error)}")
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
            raise HTTPException(status_code=500, detail=f"Text extraction failed: {str(text_error)}")
        finally:
            # Clean up temp file
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
 
        if not text or not text.strip():
            logger.error("No text could be extracted from the file using any method")
            logger.error(f"File type: {file_key.split('.')[-1].lower()}")
            
            # Provide more helpful error message
            if file_key.lower().endswith('.pdf'):
                raise HTTPException(
                    status_code=400, 
                    detail='Unable to extract text from PDF. This might be a scanned document or image-based PDF. Please ensure the PDF contains selectable text.'
                )
            else:
                raise HTTPException(
                    status_code=400, 
                    detail='No text content could be extracted from the file. Please ensure the file contains readable text content.'
                )
 
        logger.info(f"Final extracted text length: {len(text)} characters")
        logger.info(f"Text preview (first 200 chars): {text[:200]}...")
 
        # Generate quiz questions using Gemini API with better error handling
        try:
            # Check if Google API key is available
            google_api_key = os.getenv('GOOGLE_API_KEY')
            if not google_api_key:
                logger.error("GOOGLE_API_KEY environment variable is not set")
                raise HTTPException(status_code=500, detail="Google API configuration error: API key not set")
            
            questions = await generate_quiz_questions_from_text(text)
            logger.info(f"Generated {len(questions)} quiz questions")
            
            # Return quiz with metadata for scoring
            return QuizResponse(
                questions=questions,
                existing_quiz=False,
                subject=subject,
                unit=unit,
                topic=topic,
                message="New quiz generated successfully"
            )
        except Exception as quiz_error:
            logger.error(f"Quiz generation failed: {str(quiz_error)}")
            # Return a clean error message to the user
            raise HTTPException(
                status_code=500, 
                detail="Failed to generate quiz. Please try again later."
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generating quiz from file: {str(e)}")
        import traceback; traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to generate quiz from file: {str(e)}")
 
@router.get('/completed')
def check_quiz_completed(subject: str, unit: str, topic: str, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    # Clean/normalize inputs to improve matching
    cleaned_subject = clean_subject_name(subject)
    unit_input = (unit or "").strip()
    topic_input = (topic or "").strip()

    # Primary: strict completion requires xp_topic set (ensures full course flow completion)
    strict_quiz = db.query(Quiz).filter(
        Quiz.student_id == current_user.id,
        Quiz.subject == cleaned_subject,
        Quiz.unit.ilike(unit_input),
        Quiz.topic.ilike(topic_input),
        Quiz.xp_topic.isnot(None)
    ).first()

    if strict_quiz:
        return {
            "completed": True,
            "score": strict_quiz.score,
            "total_questions": 5,
            "completed_at": strict_quiz.created_at.isoformat() if strict_quiz.created_at else None,
            "xp_earned": strict_quiz.xp_topic,
            "match_type": "strict"
        }

    # Fallback: treat as completed if a matching quiz exists (case-insensitive) with either xp_topic set or a positive score
    fallback_quiz = db.query(Quiz).filter(
        Quiz.student_id == current_user.id,
        Quiz.subject == cleaned_subject,
        Quiz.unit.ilike(unit_input),
        Quiz.topic.ilike(topic_input)
    ).order_by(desc(Quiz.created_at)).first()

    if fallback_quiz:
        if (fallback_quiz.xp_topic is not None) or (fallback_quiz.score is not None and fallback_quiz.score > 0):
            return {
                "completed": True,
                "score": fallback_quiz.score,
                "total_questions": 5,
                "completed_at": fallback_quiz.created_at.isoformat() if fallback_quiz.created_at else None,
                "xp_earned": fallback_quiz.xp_topic or 0,
                "match_type": "fallback"
            }

        # Additional safeguard: if any quiz record exists for the topic, treat it as completed
        return {
            "completed": True,
            "score": fallback_quiz.score,
            "total_questions": 5,
            "completed_at": fallback_quiz.created_at.isoformat() if fallback_quiz.created_at else None,
            "xp_earned": fallback_quiz.xp_topic or 0,
            "match_type": "presence"
        }

    return {"completed": False}


@router.get('/subject-completion')
def get_subject_completion_status(
    subject: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Get completion status for all topics in a subject in a single request.
    This prevents the N+1 request problem on the frontend.
    """
    try:
        cleaned_subject = clean_subject_name(subject)
        logger.info(f"🔍 Fetching bulk completion status for subject: {cleaned_subject}")
        
        # Fetch all quizzes for this user and subject
        quizzes = db.query(Quiz).filter(
            Quiz.student_id == current_user.id,
            Quiz.subject == cleaned_subject
        ).all()
        
        # Map topic names to completion status
        completion_map = {}
        for quiz in quizzes:
            # If we already marked it true, don't overwrite with false from an older attempt
            if completion_map.get(quiz.topic):
                continue
                
            is_completed = (quiz.xp_topic is not None) or (quiz.score is not None and quiz.score > 0)
            if is_completed:
                completion_map[quiz.topic] = True
            else:
                # Even if score is 0, if the record exists, it might count as completed in some logic
                # But let's stay consistent with the /completed endpoint
                completion_map[quiz.topic] = True # Based on /completed safeguard line 973
                
        return {
            "subject": cleaned_subject,
            "completed_topics": completion_map
        }
    except Exception as e:
        logger.error(f"Error in subject-completion: {str(e)}")
        return {"subject": subject, "completed_topics": {}, "error": str(e)}


@router.get('/recent-questions')
def get_recent_quiz_questions(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Find the most recent quiz for the current user and return up to 5 questions.
    Response: {topic: [question1, question2, ...]} (only 1 topic, max 5 questions)
    """
    # Find the most recent quiz for the user across all topics
    most_recent_quiz = db.query(Quiz).filter(
        Quiz.student_id == current_user.id
    ).order_by(desc(Quiz.created_at)).first()
    
    if not most_recent_quiz:
        return {}
    
    # Get up to 5 questions from the most recent quiz
    questions = (
        db.query(QuizAttemptedQuestion)
        .filter(QuizAttemptedQuestion.quiz_id == most_recent_quiz.id)
        .order_by(QuizAttemptedQuestion.id)
        .limit(5)
        .all()
    )
    
    topic_questions = {
        most_recent_quiz.topic: [q.question_text for q in questions]
    }
    
    return topic_questions
 
@router.get('/completed-topics-count')
def get_completed_topics_count(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Get the count of unique topics the user has completed quizzes for.
    Also returns the list of completed topics for reference.
    Includes pre-assessment quiz in the count.
    """
    # Get distinct topics for which the user has completed quizzes
    topics = db.query(Quiz.topic).filter(Quiz.student_id == current_user.id).distinct().all()
   
    # Include all topics including pre-assessment
    topic_names = [topic[0] for topic in topics]
   
    # Check if pre-assessment exists
    has_preassessment = any(
        topic.lower().startswith('pre') or 'assessment' in topic.lower() 
        for topic in topic_names
    )
   
    return {
        "completed_topics_count": len(topic_names),
        "completed_topics": topic_names,
        "has_preassessment": has_preassessment,
        "meets_minimum_requirement": len(topic_names) >= 3
    }

@router.get('/completed-topics')
def get_all_completed_topics(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Get all unique topics the user has completed across all subjects.
    Optimized for bulk completion checking on the frontend.
    """
    try:
        # Get distinct topics for which the user has completed quizzes with XP or positive score
        # Using a raw SQL-like approach through SQLAlchemy for maximum speed
        topics = db.query(Quiz.topic).filter(
            Quiz.student_id == current_user.id,
            ((Quiz.xp_topic.isnot(None)) | (Quiz.score > 0))
        ).distinct().all()
        
        topic_list = [t[0] for t in topics] if topics else []
        
        return {
            "completed_topics": topic_list
        }
    except Exception as e:
        logger.error(f"Error in get_all_completed_topics: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to fetch completed topics")
 
@router.get('/performance-analysis')
def get_performance_analysis(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Analyze user's quiz performance to identify strengths and weaknesses.
    Returns performance data, strengths, weaknesses, and recommendations.
    """
    from app.models.study_plan import StudyPlanStats, StudyPlanTask
    
    # Get all quizzes for the user
    quizzes = db.query(Quiz).filter(Quiz.student_id == current_user.id).all()
    
    print(f"🔍 Performance Analysis Debug:")
    print(f"   - User ID: {current_user.id}")
    print(f"   - Total quizzes found: {len(quizzes)}")
    
    if quizzes:
        print(f"   - Sample quiz data:")
        for quiz in quizzes[:3]:
            print(f"     * Subject: {quiz.subject}, Topic: {quiz.topic}, Score: {quiz.score}")
   
    if not quizzes:
        # Get user's study plan stats for consistent metrics
        user_stats = db.query(StudyPlanStats).filter(StudyPlanStats.user_id == current_user.id).first()
        
        # Get user's participation data (historical percentage)
        user_total_tasks = db.query(StudyPlanTask).filter(
            StudyPlanTask.user_id == current_user.id,
            StudyPlanTask.date >= "2024-01-01"
        ).count()
        
        user_completed_tasks = db.query(StudyPlanTask).filter(
            StudyPlanTask.user_id == current_user.id,
            StudyPlanTask.completed == True,
            StudyPlanTask.date >= "2024-01-01"
        ).count()
        
        user_participation = int((user_completed_tasks / user_total_tasks) * 100) if user_total_tasks > 0 else 0

        # Simplified Participation (tasks completed today) logic
        today_str = date.today().strftime('%Y-%m-%d')
        # Your participation = count of today's tasks with quiz_completed == True
        your_participation = (
            db.query(func.count(StudyPlanTask.id))
            .filter(
                StudyPlanTask.user_id == current_user.id,
                StudyPlanTask.date == today_str,
                StudyPlanTask.quiz_completed == True,
            )
            .scalar()
        ) or 0

        # Cohort topper participation = max count in same year/semester for today
        user_year = current_user.year or 1
        user_semester = current_user.semester or 1
        topper_participation = (
            db.query(func.count(StudyPlanTask.id))
            .join(User, User.id == StudyPlanTask.user_id)
            .filter(
                User.year == user_year,
                User.semester == user_semester,
                StudyPlanTask.date == today_str,
                StudyPlanTask.quiz_completed == True,
            )
            .group_by(StudyPlanTask.user_id)
            .order_by(func.count(StudyPlanTask.id).desc())
            .limit(1)
            .scalar()
        )
        if topper_participation is None:
            topper_participation = your_participation
        
        # Get study hours from stats
        user_study_hours = 0
        if user_stats and user_stats.hours_this_week:
            try:
                user_study_hours = int(user_stats.hours_this_week.replace('h', ''))
            except:
                user_study_hours = 0
        
        # Calculate completion rate for users with no quizzes
        from app.api.study_plan import get_total_topics_for_year_semester
        
        # Get user's year and semester
        user_year = current_user.year or 1
        user_semester = current_user.semester or 1
        
        # Get total available topics for user's year/semester
        total_available_topics = get_total_topics_for_year_semester(user_year, user_semester)
        
        # No completed topics if no quizzes
        completed_topics = 0
        completion_rate = 0
        
        return {
            "total_quizzes": 0,
            "average_score": 0,
            "average_percentage": 0,
            "subject_performance": {},
            "topic_performance": {},
            "strengths": [],
            "weaknesses": [],
            "overall_category": "Beginner",
            "engagement_level": "Getting Started",
            # Add consistent metrics
            "study_hours": user_study_hours,
            "quiz_performance": 0,
            "participation": user_participation,
            # New fields for today's counts
            "your_participation": int(your_participation),
            "topper_participation": int(topper_participation),
            "completion_rate": completion_rate,
            "completed_topics": completed_topics,
            "total_available_topics": total_available_topics
        }
   
    # Calculate overall statistics (including all quizzes for performance analysis)
    total_quizzes = len(quizzes)
    total_score = sum(quiz.score for quiz in quizzes)
       
    average_score = total_score / total_quizzes if total_quizzes > 0 else 0
    average_percentage = (average_score / 5) * 100  # Assuming 5 questions per quiz
   
    # Group by subject-topic combination (aggregate all quizzes for same subject-topic)
    subject_performance = {}
    for quiz in quizzes:
        subject = quiz.subject
        topic = quiz.topic
        # Create a unique key for subject-topic combination
        subject_topic_key = f"{subject} - {topic}"
        
        if subject_topic_key not in subject_performance:
            subject_performance[subject_topic_key] = {
                "subject": subject,
                "topic": topic,
                "total_score": 0, 
                "total_quizzes": 0, 
                "quiz_scores": []  # Store individual scores for better calculation
            }
        
        # Add this quiz's score to the aggregation
        subject_performance[subject_topic_key]["total_score"] += quiz.score
        subject_performance[subject_topic_key]["total_quizzes"] += 1
        subject_performance[subject_topic_key]["quiz_scores"].append(quiz.score)
    
    print(f"   - Subject-topic combinations: {len(subject_performance)}")
    for key, data in list(subject_performance.items())[:3]:
        print(f"     * {key}: {data['total_quizzes']} quizzes, total score: {data['total_score']}")
   
    # Calculate subject-topic averages
    for subject_topic_key, data in subject_performance.items():
        avg_score = data["total_score"] / data["total_quizzes"]
        data["average_score"] = avg_score
        data["average_percentage"] = (avg_score / 5) * 100
        # Remove quiz_scores as it's not needed in response
        del data["quiz_scores"]
   
    # Group by topic (including all quizzes for performance analysis)
    topic_performance = {}
    for quiz in quizzes:
        topic = quiz.topic
           
        if topic not in topic_performance:
            topic_performance[topic] = {
                "total_score": 0, 
                "total_quizzes": 0, 
                "subject": quiz.subject,
                "quiz_scores": []
            }
        
        # Add this quiz's score to the aggregation
        topic_performance[topic]["total_score"] += quiz.score
        topic_performance[topic]["total_quizzes"] += 1
        topic_performance[topic]["quiz_scores"].append(quiz.score)
   
    # Calculate topic averages
    for topic, data in topic_performance.items():
        avg_score = data["total_score"] / data["total_quizzes"]
        data["average_score"] = avg_score
        data["average_percentage"] = (avg_score / 5) * 100
        # Remove quiz_scores as it's not needed in response
        del data["quiz_scores"]
   
    # Identify strengths (subject-topics with >80% average)
    strengths = []
    seen_strengths = set()  # Track seen combinations to avoid duplicates
    
    print(f"   - Analyzing strengths and weaknesses:")
    for subject_topic_key, data in subject_performance.items():
        print(f"     * {subject_topic_key}: {data['average_percentage']:.1f}%")
        if data["average_percentage"] >= 80:
            # Create a unique identifier to avoid duplicates
            strength_id = f"{data['subject']}_{data['topic']}"
            if strength_id not in seen_strengths:
                seen_strengths.add(strength_id)
                strengths.append({
                    "type": "subject_topic",
                    "name": subject_topic_key,
                    "subject": data["subject"],
                    "topic": data["topic"],
                    "percentage": round(data["average_percentage"], 1),
                    "description": f"Excellent performance in {data['subject']} - {data['topic']}"
                })
                print(f"       -> Added to strengths: {data['average_percentage']:.1f}%")
   
    for topic, data in topic_performance.items():
        if data["average_percentage"] >= 80:
            strengths.append({
                "type": "topic",
                "name": topic,
                "percentage": round(data["average_percentage"], 1),
                "description": f"Strong understanding of {topic}",
                "subject": data["subject"]
            })
   
    # Identify weaknesses (subject-topics with <60% average)
    weaknesses = []
    seen_weaknesses = set()  # Track seen combinations to avoid duplicates
    
    for subject_topic_key, data in subject_performance.items():
        if data["average_percentage"] < 60:
            # Create a unique identifier to avoid duplicates
            weakness_id = f"{data['subject']}_{data['topic']}"
            if weakness_id not in seen_weaknesses:
                seen_weaknesses.add(weakness_id)
                weaknesses.append({
                    "type": "subject_topic",
                    "name": subject_topic_key,
                    "subject": data["subject"],
                    "topic": data["topic"],
                    "percentage": round(data["average_percentage"], 1),
                    "description": f"Needs improvement in {data['subject']} - {data['topic']}"
                })
                print(f"       -> Added to weaknesses: {data['average_percentage']:.1f}%")
    
    print(f"   - Final results: {len(strengths)} strengths, {len(weaknesses)} weaknesses")
   
    for topic, data in topic_performance.items():
        if data["average_percentage"] < 60:
            weaknesses.append({
                "type": "topic",
                "name": topic,
                "percentage": round(data["average_percentage"], 1),
                "description": f"Requires more practice in {topic}",
                "subject": data["subject"]
            })
   
    # Determine overall category
    if average_percentage >= 80:
        overall_category = "Category A"
        engagement_level = "High Engagement"
    elif average_percentage >= 65:
        overall_category = "Category B"
        engagement_level = "Medium Engagement"
    else:
        overall_category = "Category C"
        engagement_level = "Low Engagement"
   
    # Get user's study plan stats for consistent metrics
    user_stats = db.query(StudyPlanStats).filter(StudyPlanStats.user_id == current_user.id).first()
    
    # Get user's participation data (historical percentage)
    user_total_tasks = db.query(StudyPlanTask).filter(
        StudyPlanTask.user_id == current_user.id,
        StudyPlanTask.date >= "2024-01-01"
    ).count()
    user_completed_tasks = db.query(StudyPlanTask).filter(
        StudyPlanTask.user_id == current_user.id,
        StudyPlanTask.completed == True,
        StudyPlanTask.date >= "2024-01-01"
    ).count()
    user_participation = int((user_completed_tasks / user_total_tasks) * 100) if user_total_tasks > 0 else 0

    # Simplified Participation (tasks completed today) logic
    today_str = date.today().strftime('%Y-%m-%d')
    # Your participation = count of today's tasks with quiz_completed == True
    your_participation = (
        db.query(func.count(StudyPlanTask.id))
        .filter(
            StudyPlanTask.user_id == current_user.id,
            StudyPlanTask.date == today_str,
            StudyPlanTask.quiz_completed == True,
        )
        .scalar()
    ) or 0

    # Cohort topper participation = max count in same year/semester for today
    user_year = current_user.year or 1
    user_semester = current_user.semester or 1
    topper_participation = (
        db.query(func.count(StudyPlanTask.id))
        .join(User, User.id == StudyPlanTask.user_id)
        .filter(
            User.year == user_year,
            User.semester == user_semester,
            StudyPlanTask.date == today_str,
            StudyPlanTask.quiz_completed == True,
        )
        .group_by(StudyPlanTask.user_id)
        .order_by(func.count(StudyPlanTask.id).desc())
        .limit(1)
        .scalar()
    )
    if topper_participation is None:
        topper_participation = your_participation
    
    # Get study hours from stats
    user_study_hours = 0
    if user_stats and user_stats.hours_this_week:
        try:
            user_study_hours = int(user_stats.hours_this_week.replace('h', ''))
        except:
            user_study_hours = 0
    
    # Calculate completion rate based on completed topics vs total available topics
    from app.api.study_plan import get_total_topics_for_year_semester
    
    # Get total available topics for user's year/semester (user_year and user_semester already defined above)
    total_available_topics = get_total_topics_for_year_semester(user_year, user_semester)
    
    # Get completed topics (unique topics with completed quizzes)
    completed_topics = db.query(Quiz.topic).filter(
        Quiz.student_id == current_user.id,
        Quiz.xp_topic.isnot(None)  # Only quizzes that were actually completed
    ).distinct().count()
    
    # Calculate completion rate (capped at 100%)
    completion_rate = min(100, int((completed_topics / total_available_topics) * 100)) if total_available_topics > 0 else 0
    
    return {
        "total_quizzes": total_quizzes,
        "average_score": round(average_score, 1),
        "average_percentage": round(average_percentage, 1),
        "subject_performance": subject_performance,
        "topic_performance": topic_performance,
        "strengths": strengths,  # All strengths
        "weaknesses": weaknesses,  # All weaknesses
        "overall_category": overall_category,
        "engagement_level": engagement_level,
        # Add consistent metrics
        "study_hours": user_study_hours,
        "quiz_performance": min(100, int(average_percentage)),
        "participation": user_participation,
        # New fields for today's counts
        "your_participation": int(your_participation),
        "topper_participation": int(topper_participation),
        "completion_rate": completion_rate,
        "completed_topics": completed_topics,
        "total_available_topics": total_available_topics
    }
 
@router.get('/analysis-data')
def get_quiz_analysis_data(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Get actual strengths and areas for improvement from quiz analysis data.
    Returns real analysis data stored in the quiz table.
    """
    try:
        logger.info(f"Fetching quiz analysis data for user {current_user.id}")
        
        # Get all quizzes with analysis data for the user
        quizzes = db.query(Quiz).filter(
            Quiz.student_id == current_user.id,
            Quiz.strengths.isnot(None),
            Quiz.weakness.isnot(None),
            Quiz.areas_to_improve.isnot(None)
        ).order_by(Quiz.created_at.desc()).all()
        
        if not quizzes:
            logger.info(f"No quiz analysis data found for user {current_user.id}")
            return {
                "strengths": [],
                "areas_for_improvement": [],
                "message": "No analysis data available yet. Complete some quizzes to see your strengths and areas for improvement."
            }
        
        # Extract unique strengths and areas for improvement
        strengths = []
        areas_for_improvement = []
        seen_strengths = set()
        seen_improvements = set()
        
        for quiz in quizzes:
            # Process strengths
            if quiz.strengths and quiz.strengths.strip() and quiz.strengths != "Analysis unavailable":
                strength_key = f"{quiz.subject}_{quiz.topic}_{quiz.strengths[:50]}"
                if strength_key not in seen_strengths:
                    seen_strengths.add(strength_key)
                    strengths.append({
                        "name": f"{quiz.subject} - {quiz.topic}",
                        "description": quiz.strengths,
                        "score": round((quiz.score / 5) * 100),  # Convert to percentage
                        "subject": quiz.subject,
                        "topic": quiz.topic,
                        "quiz_id": quiz.id,
                        "created_at": quiz.created_at.isoformat() if quiz.created_at else None
                    })
            
            # Process areas for improvement
            if quiz.areas_to_improve and quiz.areas_to_improve.strip() and quiz.areas_to_improve != "Analysis unavailable":
                improvement_key = f"{quiz.subject}_{quiz.topic}_{quiz.areas_to_improve[:50]}"
                if improvement_key not in seen_improvements:
                    seen_improvements.add(improvement_key)
                    areas_for_improvement.append({
                        "name": f"{quiz.subject} - {quiz.topic}",
                        "description": quiz.areas_to_improve,
                        "score": round((quiz.score / 5) * 100),  # Convert to percentage
                        "subject": quiz.subject,
                        "topic": quiz.topic,
                        "quiz_id": quiz.id,
                        "created_at": quiz.created_at.isoformat() if quiz.created_at else None
                    })
        
        # Sort by score (descending for strengths, ascending for improvements)
        strengths.sort(key=lambda x: x["score"], reverse=True)
        areas_for_improvement.sort(key=lambda x: x["score"])
        
        logger.info(f"Found {len(strengths)} strengths and {len(areas_for_improvement)} areas for improvement for user {current_user.id}")
        
        return {
            "strengths": strengths,  # All strengths (removed [:5] limitation)
            "areas_for_improvement": areas_for_improvement,  # All areas for improvement (removed [:5] limitation)
            "total_quizzes_analyzed": len(quizzes)
        }
        
    except Exception as e:
        logger.error(f"Error fetching quiz analysis data for user {current_user.id}: {str(e)}")
        return {
            "strengths": [],
            "areas_for_improvement": [],
            "error": "Failed to fetch analysis data",
            "message": "Unable to retrieve analysis data at this time."
        }

@router.get('/search-by-subject')
def search_quizzes_by_subject(
    subject_name: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Search for quiz details by subject name for the current user.
    Returns all quizzes for the specified subject with detailed information.
    """
    try:
        logger.info(f"Searching quizzes for subject: {subject_name} for user {current_user.id}")
        
        # Clean the subject name to match how it's stored in the database
        cleaned_subject_name = clean_subject_name(subject_name)
        
        # Search for quizzes by subject name (case-insensitive, exact match preferred)
        # First try exact match, then partial match
        quizzes = db.query(Quiz).filter(
            Quiz.student_id == current_user.id,
            Quiz.subject.ilike(f"%{cleaned_subject_name}%")
        ).order_by(Quiz.created_at.desc()).all()
        
        # If no results with partial match, try exact match
        if not quizzes:
            quizzes = db.query(Quiz).filter(
                Quiz.student_id == current_user.id,
                Quiz.subject == cleaned_subject_name
            ).order_by(Quiz.created_at.desc()).all()
        
        if not quizzes:
            return {
                "subject_name": subject_name,
                "total_quizzes": 0,
                "quizzes": [],
                "message": f"No quizzes found for subject: {subject_name}"
            }
        
        # Get detailed information for each quiz
        quiz_details = []
        for quiz in quizzes:
            # Get attempted questions for this quiz
            attempted_questions = db.query(QuizAttemptedQuestion).filter(
                QuizAttemptedQuestion.quiz_id == quiz.id
            ).all()
            
            quiz_detail = {
                "quiz_id": quiz.id,
                "subject": quiz.subject,
                "unit": quiz.unit,
                "topic": quiz.topic,
                "score": quiz.score,
                "total_questions": len(attempted_questions),
                "percentage": (quiz.score / len(attempted_questions) * 100) if attempted_questions else 0,
                "year": quiz.year,
                "semester": quiz.semester,
                "completed_at": quiz.created_at.isoformat() if quiz.created_at else None,
                "questions": [q.question_text for q in attempted_questions]
            }
            quiz_details.append(quiz_detail)
        
        # Calculate summary statistics
        total_quizzes = len(quizzes)
        average_score = sum(q.score for q in quizzes) / total_quizzes if total_quizzes > 0 else 0
        average_percentage = sum((q.score / len(db.query(QuizAttemptedQuestion).filter(QuizAttemptedQuestion.quiz_id == q.id).all()) * 100) for q in quizzes) / total_quizzes if total_quizzes > 0 else 0
        
        # Get unique topics and units
        unique_topics = list(set(q.topic for q in quizzes))
        unique_units = list(set(q.unit for q in quizzes))
        
        return {
            "subject_name": subject_name,
            "total_quizzes": total_quizzes,
            "average_score": round(average_score, 2),
            "average_percentage": round(average_percentage, 2),
            "unique_topics": unique_topics,
            "unique_units": unique_units,
            "quizzes": quiz_details,
            "summary": {
                "best_score": max(q.score for q in quizzes) if quizzes else 0,
                "worst_score": min(q.score for q in quizzes) if quizzes else 0,
                "total_questions_attempted": sum(len(db.query(QuizAttemptedQuestion).filter(QuizAttemptedQuestion.quiz_id == q.id).all()) for q in quizzes),
                "first_quiz_date": min(q.created_at for q in quizzes).isoformat() if quizzes else None,
                "last_quiz_date": max(q.created_at for q in quizzes).isoformat() if quizzes else None
            }
        }
        
    except Exception as e:
        logger.error(f"Error searching quizzes by subject: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to search quizzes: {str(e)}")

@router.get('/study-plan')
def generate_study_plan(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Generate a personalized study plan based on quiz performance analysis.
    """
    # Get performance analysis
    performance_analysis = get_performance_analysis(db, current_user)
   
    if performance_analysis["total_quizzes"] == 0:
        return {
            "plan_type": "Getting Started",
            "description": "Complete more quizzes to generate a personalized study plan",
            "tasks": [
                {
                    "id": 1,
                    "type": "quiz",
                    "title": "Take Your First Quiz",
                    "description": "Start with any topic to establish your baseline",
                    "priority": "high",
                    "estimated_time": "15 min",
                    "status": "pending"
                }
            ],
            "focus_areas": [],
            "recommendations": [
                "Start with topics you're most interested in",
                "Take quizzes regularly to track your progress",
                "Review materials before taking quizzes"
            ]
        }
   
    # Generate tasks based on weaknesses
    tasks = []
    task_id = 1
   
    # Add tasks for weak areas (excluding pre-assessment topics)
    for weakness in performance_analysis["weaknesses"]:
        if weakness["type"] == "topic" and not (weakness["name"].lower().startswith('pre') or 'assessment' in weakness["name"].lower()):
            tasks.append({
                "id": task_id,
                "type": "review",
                "title": f"Review {weakness['name']}",
                "description": f"Focus on improving {weakness['name']} (Current: {weakness['percentage']}%)",
                "priority": "high",
                "estimated_time": "30 min",
                "status": "pending",
                "subject": weakness.get("subject", "General")
            })
            task_id += 1
           
            tasks.append({
                "id": task_id,
                "type": "quiz",
                "title": f"Practice Quiz: {weakness['name']}",
                "description": f"Test your improved knowledge in {weakness['name']}",
                "priority": "medium",
                "estimated_time": "15 min",
                "status": "pending",
                "subject": weakness.get("subject", "General")
            })
            task_id += 1
   
    # Add reinforcement tasks for strengths (excluding pre-assessment topics)
    for strength in performance_analysis["strengths"][:2]:  # Top 2 strengths
        if strength["type"] == "topic" and not (strength["name"].lower().startswith('pre') or 'assessment' in strength["name"].lower()):
            tasks.append({
                "id": task_id,
                "type": "advanced",
                "title": f"Advanced Practice: {strength['name']}",
                "description": f"Challenge yourself with advanced concepts in {strength['name']}",
                "priority": "low",
                "estimated_time": "25 min",
                "status": "pending",
                "subject": strength.get("subject", "General")
            })
            task_id += 1
   
    # Generate plan description based on performance
    if performance_analysis["average_percentage"] >= 80:
        plan_type = "Advanced Mastery"
        description = "Focus on advanced concepts and challenge yourself with complex problems"
    elif performance_analysis["average_percentage"] >= 65:
        plan_type = "Skill Enhancement"
        description = "Strengthen weak areas while maintaining your current knowledge level"
    else:
        plan_type = "Foundation Building"
        description = "Build a strong foundation by focusing on fundamental concepts"
   
    # Generate focus areas
    focus_areas = []
    weak_subjects = set()
    for weakness in performance_analysis["weaknesses"]:
        if weakness["type"] == "subject":
            weak_subjects.add(weakness["name"])
   
    for subject in weak_subjects:
        focus_areas.append({
            "subject": subject,
            "priority": "high",
            "reason": f"Below average performance in {subject}"
        })
   
    # Generate recommendations
    recommendations = []
    if performance_analysis["average_percentage"] < 60:
        recommendations.extend([
            "Focus on fundamental concepts before attempting advanced topics",
            "Take more time to review materials before quizzes",
            "Practice regularly to improve retention"
        ])
    elif performance_analysis["average_percentage"] < 80:
        recommendations.extend([
            "Identify and work on specific weak areas",
            "Balance review time between weak and strong subjects",
            "Set aside dedicated time for challenging topics"
        ])
    else:
        recommendations.extend([
            "Challenge yourself with advanced problems",
            "Help others to reinforce your knowledge",
            "Explore related topics to broaden your understanding"
        ])
   
        return {
            "plan_type": plan_type,
            "description": description,
            "tasks": tasks[:6],  # Limit to 6 tasks
            "focus_areas": focus_areas,
            "recommendations": recommendations,
            "performance_summary": {
                "total_quizzes": performance_analysis["total_quizzes"],
                "average_percentage": performance_analysis["average_percentage"],
                "category": performance_analysis["overall_category"],
                "engagement": performance_analysis["engagement_level"]
            }
        }

@router.get('/analysis/{quiz_id}')
def get_quiz_analysis(
    quiz_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Get detailed analysis for a specific quiz.
    """
    try:
        logger.info(f"Fetching quiz analysis for quiz {quiz_id}, user {current_user.id}")
        
        # Get the quiz record
        quiz = db.query(Quiz).filter(
            Quiz.id == quiz_id,
            Quiz.student_id == current_user.id
        ).first()
        
        if not quiz:
            raise HTTPException(status_code=404, detail="Quiz not found")
        
        # Get attempted questions for this quiz
        attempted_questions = db.query(QuizAttemptedQuestion).filter(
            QuizAttemptedQuestion.quiz_id == quiz.id
        ).all()
        
        return {
            "quiz_id": quiz.id,
            "subject": quiz.subject,
            "unit": quiz.unit,
            "topic": quiz.topic,
            "score": quiz.score,
            "total_questions": len(attempted_questions),
            "percentage": (quiz.score / len(attempted_questions) * 100) if attempted_questions else 0,
            "strengths": quiz.strengths,
            "weakness": quiz.weakness,
            "areas_to_improve": quiz.areas_to_improve,
            "year": quiz.year,
            "semester": quiz.semester,
            "completed_at": quiz.created_at.isoformat() if quiz.created_at else None,
            "questions": [q.question_text for q in attempted_questions]
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching quiz analysis: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
 


